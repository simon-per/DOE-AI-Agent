"""
Stage 4: Generate tailored cover letters (PDF + DOCX) for high-scoring jobs.
Primary: OpenRouter (Qwen3 235B) - reliable, paid (low cost, no rate limits)
Fallback: Google AI Studio (Gemini 3 Flash) - free but rate-limited

Auto-detects job language (DE/EN/IT) and writes cover letter accordingly.
Only generates for jobs scoring >= 6 (configurable via --min-score).

Input:  .tmp/scored_jobs.json
Output: .tmp/applications/{company}_{title}/Cover_Letter_Simon_Oberpertinger_Mair_{company}.pdf + .docx

Usage:
    python execution/generate_cover_letter.py
    python execution/generate_cover_letter.py --limit 2              # test with 2 jobs
    python execution/generate_cover_letter.py --min-score 8           # only top matches
    python execution/generate_cover_letter.py --reset-checkpoint      # re-generate all
"""

import argparse
import json
import logging
import os
import random
import re
import sys
import time
from datetime import datetime
from pathlib import Path

import html as html_module

import requests
from dotenv import load_dotenv
from fpdf import FPDF

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))
load_dotenv(PROJECT_ROOT / ".env")

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
log = logging.getLogger(__name__)

# Suppress noisy font subsetting logs from fpdf2/fontTools
logging.getLogger("fontTools.subset").setLevel(logging.WARNING)

TMP_DIR = PROJECT_ROOT / ".tmp"
INPUT_FILE = TMP_DIR / "scored_jobs.json"
OUTPUT_DIR = TMP_DIR / "applications"
COMPANY_CACHE_FILE = TMP_DIR / "company_cache.json"
CHECKPOINT_FILE = TMP_DIR / "cover_letter_checkpoint.json"

DEFAULT_MIN_SCORE = 6

# Shared modules
from execution.llm_client import call_llm as _call_llm_shared
from execution.language_detect import detect_language
from execution.utils import (
    sanitize_filename,
    strip_legal_suffixes,
    clean_job_title,
    generate_job_id,
    enforce_revops_subtitle,
    find_application_folder,
    has_cover_letter_outputs,
)

# Profile loaded from ABOUTME.md (single source of truth)
from execution.profile_loader import load_profile as _load_profile

# Calibration: Qwen3 undershoots ~40-50 words in German. Prompt target 260-300 yields actual ~220-260 output.
# Do NOT put this note inside the prompt — the model uses it as permission to write shorter.
COVER_LETTER_PROMPT = """You are writing a professional cover letter for a Swiss job application.

CANDIDATE PROFILE:
{profile}

JOB DETAILS:
Title: {title}
Company: {company}
Location: {location}
Key Matches: {key_matches}
Key Gaps: {key_gaps}

COMPANY CONTEXT:
{company_context}

JOB DESCRIPTION (excerpt):
{description}

STYLE EXAMPLE (match this tone, directness, and structure — do NOT copy the content):
\"Als führender Anbieter integrierter Logistiklösungen für die Schweizer Pharmaindustrie hat MedLog AG genau das Umfeld, in dem meine dreijährige Erfahrung in Datenanalyse und Prozessdigitalisierung den grössten Mehrwert liefert. Ihr Fokus auf temperaturgeführte Lieferketten und die regulatorischen Anforderungen decken sich mit meinem Arbeitsalltag.

In meiner aktuellen Rolle habe ich über 150 automatisierte Reports in Power BI aufgebaut, die Lagerbestandsgenauigkeit von 84% auf 97% gesteigert und SAP-basierte Bestellprozesse für 12 Standorte standardisiert. Als technischer Ansprechpartner für ein 20-köpfiges Operations-Team übersetze ich täglich fachliche Anforderungen in datenbasierte Lösungen — von der Kennzahlenanalyse bis zur Prozessautomation.

Mein technisches Profil stütze ich durch den Google Data Analytics Professional Certificate und fortlaufende Weiterbildung in SQL und API-Integrationen. Die Position als Data Analyst Supply Chain bietet mir die ideale Möglichkeit, diese Kompetenzen in einem regulierten Umfeld weiterzuentwickeln.

Für ein persönliches Gespräch stehe ich gerne zur Verfügung — meine Kündigungsfrist beträgt 45 Tage.\"

(The style example above is ~150 words for brevity — your output must be 260-300 words, i.e. notably longer than the example.)

ANTI-HALLUCINATION RULES (CRITICAL — violations cause interview failures):
- You may ONLY use numbers and metrics that appear in the CANDIDATE PROFILE above.
- METRIC SELECTION (MANDATORY — use EXACTLY 3 or 4, NEVER all 6):
  The candidate has 7 quantitative achievements. You MUST pick exactly 3 or 4 that are MOST relevant to THIS specific job. Leave the rest out entirely — do NOT mention them even without numbers.
  Available metrics:
  A) 400,000+ CRM records managed
  B) 1,000+ bulk data uploads executed
  C) 20% data integrity improvement
  D) 100+ CPQ product attributes configured
  E) 200+ SAP material masters maintained
  F) 30-person sales team supported
  G) 10 Power BI dashboards built
  HARD RULE: Count the metrics you used. If you used 5 or more, you MUST rewrite and drop extras. Using all metrics makes every letter identical — this is the #1 reason cover letters get rejected as template-generated.
- Do NOT invent percentages, revenue figures, team sizes, time savings, or any numbers not listed above.
- Do NOT claim skills not listed in the profile.

BANNED PHRASES (never use these overused clichés):
- "Mit grossem Interesse habe ich Ihre Stellenausschreibung gelesen"
- "Hiermit bewerbe ich mich auf die Stelle als"
- "Ihre Stellenausschreibung hat mich sofort angesprochen" / "begeistert"
- "Auf der Suche nach einer neuen Herausforderung"
- "Gerne möchte ich mich Ihnen vorstellen"
- "Als motivierter und teamfähiger"
- "Ich bin überzeugt, dass ich die richtige Person bin"
- Do NOT start 3+ consecutive sentences with "Ich"

BANNED AI-GIVEAWAY WORDS (never use — recruiters use these to detect AI):
- English: delve, utilize, leverage, moreover, furthermore, compelling, robust, streamline, facilitate, paramount, embark, navigate, multifaceted, transformative, foster, elevate, realm, bolster, forge, unleash, results-driven, operational excellence, proven expertise, proven track record
- German: darüber hinaus, zweifellos, massgeblich, wegweisend, richtungsweisend, bahnbrechend, ganzheitlich, ergebnisorientiert

INSTRUCTIONS:
- Write in {language_name} ({language_code}). If not German, adapt the style example's directness to that language.
- Length: 260-300 words. This is the ideal range — letters under 220 words appear too brief and reduce callbacks. Do NOT write less than 260 words. Count carefully.
  When using fewer metrics (3-4), use the freed space for deeper qualitative descriptions — explain HOW you did things, not just WHAT. This keeps the letter substantive.
- Tone: Direct, Swiss-professional, modest but confident. NOT aggressive American self-promotion. Every claim backed by a specific number from the profile or a tool name.
- Sentence length variation (prevents AI detection):
  Mix short (5-12 words) and longer (18-30 words) sentences naturally. Avoid uniform sentence length.
  GOOD: "Daten treiben den Vertrieb. In meiner Rolle habe ich ueber 400.000 CRM-Datensaetze verwaltet und die Integritaet um 20 Prozent gesteigert — eine Grundlage fuer automatisierte Kampagnensteuerung. Durch gezielte Segmentierung und automatisierte E-Mail-Journeys konnte ich Vertriebsprozesse beschleunigen, die vorher manuell liefen."
  BAD: "Ich habe CRM-Datensaetze verwaltet. Ich habe Massenuploads durchgefuehrt. Ich habe die Datenqualitaet verbessert." (all same length = AI-detected)
  NEVER insert meta-commentary about writing style (e.g. "Short sentences work.", "Clarity matters."). Every sentence must be substantive content.
- Structure (Problem-Solution format — highest conversion rate):
  1. OPENING (2-3 sentences, ~60 words): Identify a specific challenge or goal the company faces (from COMPANY CONTEXT or JOB DESCRIPTION), then connect YOUR experience as the solution. Show you researched them.
  2. BODY (4-5 sentences, ~140 words): Your 2-3 strongest concrete examples with REAL numbers from your profile that match THIS job. Mirror exact terminology from the job description (e.g. "Kundenstammdatenpflege" not "CRM data management", "Offertenwesen" not "quotation management").
  3. CLOSING (1-2 sentences, ~50 words): Brief forward look + availability (45 days notice). No filler.
- Do NOT include addresses, date, subject line, greeting, or signature (added by PDF formatter)
- Start directly with the opening paragraph
- Focus on KEY MATCHES: {key_matches}
- If there are skill gaps, frame them briefly as growth opportunities (one sentence max)
- Do NOT mention lack of a university degree
- Swiss German sign-off uses "Freundliche Grüsse" (NOT "Mit freundlichen Grüssen")

CRITICAL: Return ONLY valid JSON:
{{"text": "The cover letter text (paragraphs separated by \\n\\n)", "subtitle": "Revenue Operations Specialist | <3-5 skill keywords matching this job's stack>"}}

SUBTITLE RULES (must follow exactly):
- Format: "Revenue Operations Specialist | <suffix>"
- The prefix "Revenue Operations Specialist" is fixed — no translation, no synonyms (NOT "Spezialist", NOT "Sales Operations Specialist", NOT the literal job title).
- The suffix is 3-5 short skill keywords from THIS job's tech stack.
- Use ONLY " • " (bullet with surrounding spaces) as the separator between keywords.
- Each keyword is ONE concept. NEVER join two distinct tools with a space (e.g. "SAP CPQ" is WRONG because SAP and CPQ are two different products in this profile — write "SAP • CPQ" or "SAP S/4HANA • CPQ" instead). The same applies to any other pair like "Power BI Excel" or "Dynamics SAP".
- Prefer generic category names (CRM, BI, CPQ, Automation, Analytics) for visual cleanness. EXCEPTION: keep specific product names when they are recognized industry-standard keywords that ATS systems explicitly filter for — specifically "Power BI", "Excel", "SAP", "Salesforce", "Dynamics". For this candidate, "Power BI" should always be written in full (not "BI") because it is a Microsoft-certified credential and a high-ROI brand keyword.
- Examples: "CRM • Power BI • Automation", "CRM • Power BI • CPQ • Automation", "SAP S/4HANA • CPQ • Process Design", "Dynamics 365 • Power BI • SQL"."""


# ---------------------------------------------------------------------------
# LLM calling — delegates to shared llm_client
# ---------------------------------------------------------------------------

def _call_llm(openrouter_key: str | None, gemini_key: str | None, prompt: str, temperature: float = 0.3) -> tuple[str, str]:
    """Call LLM and return (response_text, provider_name). Uses max_tokens=1000 for cover letters."""
    return _call_llm_shared(openrouter_key, gemini_key, prompt, temperature=temperature, max_tokens=1000)


def _get_profile_for_provider(openrouter_key: str | None) -> str:
    """Return anonymized profile if using OpenRouter, full otherwise."""
    profile = _load_profile()
    if openrouter_key:
        return profile.profile_anonymous()
    return profile.profile_full()


# ---------------------------------------------------------------------------
# Company research (web fetch + LLM extract + cache)
# ---------------------------------------------------------------------------

def _load_company_cache() -> dict:
    """Load cached company research results."""
    if COMPANY_CACHE_FILE.exists():
        try:
            with open(COMPANY_CACHE_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except (json.JSONDecodeError, OSError):
            return {}
    return {}


def _save_company_cache(cache: dict):
    """Persist company research cache to disk."""
    COMPANY_CACHE_FILE.parent.mkdir(parents=True, exist_ok=True)
    with open(COMPANY_CACHE_FILE, "w", encoding="utf-8") as f:
        json.dump(cache, f, ensure_ascii=False, indent=2)


def _guess_company_urls(company: str) -> list[str]:
    """Generate likely website URLs for a Swiss company."""
    clean = strip_legal_suffixes(company)

    if not clean or len(clean) < 2:
        return []

    words = clean.split()

    # Build candidate domain names
    candidates = []
    # First word only: "Graphax" — most common for Swiss companies
    first = words[0].lower()
    # Full hyphenated: "ergon-informatik"
    full_hyphen = "-".join(w.lower() for w in words)
    # Full concatenated: "ergoninformatik"
    full_concat = "".join(w.lower() for w in words)

    for name in [first, full_hyphen, full_concat]:
        # Replace umlauts for URLs
        name = (name.replace('ä', 'ae').replace('ö', 'oe')
                .replace('ü', 'ue').replace('ß', 'ss'))
        name = re.sub(r'[^a-z0-9-]', '', name)
        if name and name not in [c for c, _ in candidates]:
            candidates.append((name, None))

    urls = []
    seen = set()
    for name, _ in candidates:
        for tld in ['.ch', '.com', '.de']:
            url = f"https://www.{name}{tld}"
            if url not in seen:
                urls.append(url)
                seen.add(url)

    return urls[:9]  # max 9 attempts


def _fetch_webpage_text(url: str, max_chars: int = 3000) -> str | None:
    """Fetch a webpage and extract plain text content."""
    try:
        resp = requests.get(url, timeout=8, headers={
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
        }, allow_redirects=True)
        if resp.status_code != 200:
            return None
        # Skip non-HTML responses
        content_type = resp.headers.get("Content-Type", "")
        if "text/html" not in content_type:
            return None

        text = resp.text
        # Remove script and style blocks
        text = re.sub(r'<script[^>]*>.*?</script>', '', text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL | re.IGNORECASE)
        # Strip HTML tags
        text = re.sub(r'<[^>]+>', ' ', text)
        # Decode HTML entities
        text = html_module.unescape(text)
        # Normalize whitespace
        text = re.sub(r'\s+', ' ', text).strip()

        return text[:max_chars] if len(text) > 50 else None
    except Exception:
        return None


def _extract_about_from_description(description: str) -> str:
    """Extract 'About us' section or company intro from job description."""
    if not description or len(description) < 50:
        return ""

    # Look for explicit "About" section headers
    about_headers = (
        r'(?:Über uns|Wer wir sind|Unser Unternehmen|Das Unternehmen|'
        r'About us|About the company|About|Who we are|'
        r'Chi siamo|L\'azienda)'
    )
    next_section = (
        r'(?:Aufgaben|Tasks|Anforderungen|Requirements|Ihre Aufgaben|Your Tasks|'
        r'Was wir bieten|What we offer|Dein Profil|Ihr Profil|Your Profile|'
        r'Deine Aufgaben|Ihre Vorteile|Benefits|Qualifikationen|Qualifications|'
        r'Wir bieten|Stellenbeschreibung|Das bringst du mit|Was du mitbringst|'
        r'Deine Verantwortung|Ihre Verantwortung|Das erwartet dich|Unser Angebot)'
    )
    pattern = rf'{about_headers}\s*[:\n]+(.*?)(?=\n\s*{next_section}|\Z)'
    match = re.search(pattern, description, re.IGNORECASE | re.DOTALL)
    if match:
        about = match.group(1).strip()
        if len(about) > 50:
            return about[:500]

    # Fallback: first 2 sentences often describe the company
    sentences = re.split(r'(?<=[.!?])\s+', description.strip()[:600])
    if len(sentences) >= 2 and len(sentences[0]) > 30:
        return ". ".join(sentences[:2]).strip()

    return ""


def _extract_company_domain_from_job_url(job_url: str) -> str | None:
    """If the job URL is on the company's own domain (not a job board), use it."""
    JOB_BOARDS = {
        "indeed.com", "linkedin.com", "glassdoor.com", "jobs.ch",
        "jobup.ch", "jobcloud.ch", "google.com", "monster.ch",
        "indeed.ch", "glassdoor.ch", "xing.com", "stepstone.ch",
        "karriere.at", "huzzle.com",
    }
    if not job_url:
        return None
    try:
        from urllib.parse import urlparse
        parsed = urlparse(job_url)
        domain = parsed.netloc.replace("www.", "").lower()
        if domain and not any(board in domain for board in JOB_BOARDS):
            return f"https://{parsed.netloc}"
    except Exception:
        pass
    return None


def _validate_website_for_company(company_name: str, web_text: str) -> bool:
    """Check that the scraped website actually belongs to the target company.

    Strict matching: for short names (1-3 key words), ALL words must appear.
    This prevents e.g. 'KOWI Assurance AG' matching kowi.ch (fashion brand).
    """
    clean = strip_legal_suffixes(company_name)
    key_words = [w.lower() for w in clean.split() if len(w) > 2]

    if not key_words:
        return True  # Can't validate, allow it

    text_lower = web_text.lower()
    # Use word-boundary matching to prevent "SAP" matching "ASAP", etc.
    matches = sum(1 for w in key_words if re.search(rf'\b{re.escape(w)}\b', text_lower))

    # Short names (1-3 words): require ALL key words to match
    # Longer names (4+): require at least 75% to match
    if len(key_words) <= 3:
        return matches == len(key_words)
    return matches >= max(1, int(len(key_words) * 0.75))


def _research_company(
    company: str,
    description: str,
    openrouter_key: str | None,
    gemini_key: str | None,
    job_url: str = "",
    job_title: str = "",
) -> str:
    """Research company and return a brief summary for cover letter personalization.

    Primary: Uses web_search.research_company() (Google Search via SerpAPI).
    Fallback: Legacy URL-guessing + job description extraction.
    """
    # Try the new web search module (Google Search via SerpAPI)
    try:
        from execution.web_search import research_company as ws_research
        result = ws_research(
            company_name=company,
            location="Switzerland",
            job_description=description,
            job_url=job_url,
            job_title=job_title,
            openrouter_key=openrouter_key,
            gemini_key=gemini_key,
        )
        desc = result.get("description", "")
        source = result.get("source", "unknown")
        confidence = result.get("confidence", 0)
        website = result.get("website")
        if website:
            log.info(f"  Company website: {website} (source: {source}, confidence: {confidence:.0%})")
        # Don't use low-confidence data — risk of wrong company info in cover letter
        if confidence < 0.5:
            log.warning(f"  Low confidence ({confidence:.0%}) for '{company}' — discarding company context to avoid wrong info")
            return ""
        return desc
    except ImportError:
        log.warning("  web_search module not available, using legacy company research")
    except Exception as e:
        log.warning(f"  web_search failed ({e}), falling back to legacy company research")

    # Legacy fallback: cache + URL guessing + job description
    cache = _load_company_cache()
    cache_key = company.strip().lower()
    if cache_key in cache:
        cached = cache[cache_key]
        # Handle both old (string) and new (dict) cache formats
        if isinstance(cached, dict):
            cached = cached.get("description", "")
        if cached:
            log.info(f"  Company info (cached): {cached[:80]}...")
        else:
            log.info("  Company info (cached): none available")
        return cached

    about_text = _extract_about_from_description(description)

    web_text = None
    company_domain = _extract_company_domain_from_job_url(job_url)
    if company_domain:
        web_text = _fetch_webpage_text(company_domain)
        if web_text:
            if _validate_website_for_company(company, web_text):
                log.info(f"  Company website (from job URL): {company_domain}")
            else:
                log.info(f"  Job URL domain {company_domain} doesn't match company '{company}', skipping")
                web_text = None

    if not web_text:
        urls = _guess_company_urls(company)
        for url in urls:
            fetched = _fetch_webpage_text(url)
            if fetched:
                if _validate_website_for_company(company, fetched):
                    web_text = fetched
                    log.info(f"  Company website found: {url}")
                    break
                else:
                    log.info(f"  Website {url} doesn't match '{company}', skipping")

    context_parts = []
    if about_text:
        context_parts.append(f"From job posting:\n{about_text}")
    if web_text:
        context_parts.append(f"From company website:\n{web_text}")

    if not context_parts:
        log.info(f"  No company info found for {company}")
        cache[cache_key] = ""
        _save_company_cache(cache)
        return ""

    context = "\n\n".join(context_parts)
    prompt = (
        f"Extract 2-3 key facts about {company} that a job applicant could reference "
        f"in a cover letter to show they researched the company.\n"
        f"Focus on: what the company does (products/services), their industry or market "
        f"position, and any notable values or achievements.\n"
        f"Return ONLY a brief 1-2 sentence summary. Example:\n"
        f"\"Graphax is Switzerland's leading independent provider of IT services and "
        f"digital printing solutions, serving businesses across the DACH region.\"\n\n"
        f"Company info:\n{context[:2000]}"
    )

    try:
        result, provider = _call_llm(openrouter_key, gemini_key, prompt)
        result = result.strip().strip('"').strip("'")
        if len(result) > 20 and len(result) < 500:
            log.info(f"  Company research ({provider}): {result[:100]}...")
            cache[cache_key] = result
            _save_company_cache(cache)
            return result
    except Exception as e:
        log.warning(f"  Company research LLM failed: {e}")

    fallback = about_text[:200] if about_text else ""
    cache[cache_key] = fallback
    _save_company_cache(cache)
    if fallback:
        log.info(f"  Company info (description fallback): {fallback[:80]}...")
    return fallback


    # Language detection consolidated into execution/language_detect.py
    # detect_language is imported at top of file from execution.language_detect


# ---------------------------------------------------------------------------
# PDF generation
# ---------------------------------------------------------------------------

def _clean_title_for_subject(title: str) -> str:
    """Clean raw job title for use in cover letter subject line."""
    return clean_job_title(title)


    # sanitize_filename consolidated into execution/utils.py
    # sanitize_filename is imported at top of file from execution.utils


class CoverLetterPDF(FPDF):
    """Custom PDF with professional cover letter layout."""

    def __init__(self, language_code: str = "de", subtitle: str = "Revenue Operations Specialist | CRM • SAP • Power BI"):
        super().__init__()
        self.language_code = language_code
        self.subtitle = subtitle
        self._setup_fonts()

    def _setup_fonts(self):
        """Add Unicode font for proper umlaut/accent support (cross-platform)."""
        import platform
        system = platform.system()
        if system == "Windows":
            font_dir = Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts"
        elif system == "Darwin":
            font_dir = Path("/Library/Fonts")
        else:
            font_dir = Path("/usr/share/fonts/truetype/liberation")
        # Try Arial first, then Liberation Sans (Linux equivalent)
        arial_map = {"": "arial.ttf", "B": "arialbd.ttf", "I": "ariali.ttf"}
        liberation_map = {"": "LiberationSans-Regular.ttf", "B": "LiberationSans-Bold.ttf", "I": "LiberationSans-Italic.ttf"}
        for style, filename in arial_map.items():
            font_path = font_dir / filename
            if not font_path.exists() and system == "Linux":
                font_path = font_dir / liberation_map[style]
            if font_path.exists():
                self.add_font("Arial", style, str(font_path))
            else:
                log.warning(f"Font not found: {font_path} — PDF may lack proper Unicode support")

    def header(self):
        """Add candidate header to each page."""
        self.set_font("Arial", "B", 14)
        _p = _load_profile().personal
        self.cell(0, 8, _p.name, new_x="LMARGIN", new_y="NEXT")

        self.set_font("Arial", "", 9)
        self.set_text_color(100, 100, 100)
        self.cell(0, 5, self.subtitle, new_x="LMARGIN", new_y="NEXT")
        self.cell(0, 5, f"{_p.email} | {_p.phone} | {_p.location}", new_x="LMARGIN", new_y="NEXT")
        self.set_text_color(0, 0, 0)

        # Separator line
        self.ln(3)
        self.set_draw_color(200, 200, 200)
        self.line(10, self.get_y(), 200, self.get_y())
        self.ln(8)

    def footer(self):
        """Add page number to footer."""
        self.set_y(-15)
        self.set_font("Arial", "I", 8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 10, f"Page {self.page_no()}", align="C")


    # Contact extraction consolidated into execution/extract_contacts.py
from execution.extract_contacts import extract_contact_person as _extract_contact, detect_honorific


def _last_name(full_name: str) -> str:
    """Best-effort lastname extraction. 'Frau Anna Müller' → 'Müller', 'Dr. Anna Müller' → 'Müller'.

    Strips honorifics + academic titles, then takes the last whitespace-separated token.
    """
    if not full_name:
        return ""
    # Strip honorifics + common titles from the front
    cleaned = re.sub(
        r"^\s*(Frau|Herr|Ms\.?|Mr\.?|Mrs\.?|Sig\.?(?:ra)?|Dott\.?(?:ssa)?|Dr\.?|Prof\.?)\s+",
        "",
        full_name.strip(),
        flags=re.IGNORECASE,
    )
    parts = cleaned.split()
    return parts[-1] if parts else ""


def _greeting_for_contact(
    contact_name: str | None,
    language_code: str,
    gender_hint: str | None = None,
) -> str:
    """Generate appropriate greeting based on whether we have a contact name + gender.

    Backwards-compat: callers passing only (contact_name, language_code) — the historical
    signature — get exactly the same output as before. The new `gender_hint` opt-in
    upgrades the salutation to formal "Sehr geehrte Frau Müller," when known.
    """
    if not contact_name:
        fallback = {
            "de": "Sehr geehrte Damen und Herren,",
            "en": "Dear Hiring Manager,",
            "it": "Gentili Signore e Signori,",
        }
        return fallback.get(language_code, fallback["de"])

    # If gender is known (Frau / Herr / Ms / Mr), use the formal "Sehr geehrte Frau {Lastname}"
    # form which recruiters in DACH expect from a serious application.
    last = _last_name(contact_name)
    hint_lower = (gender_hint or "").strip().lower()
    if last and hint_lower in ("frau", "ms", "mrs"):
        formal = {
            "de": f"Sehr geehrte Frau {last},",
            "en": f"Dear Ms. {last},",
            "it": f"Gentile Sig.ra {last},",
        }
        return formal.get(language_code, formal["de"])
    if last and hint_lower in ("herr", "mr"):
        formal = {
            "de": f"Sehr geehrter Herr {last},",
            "en": f"Dear Mr. {last},",
            "it": f"Gentile Sig. {last},",
        }
        return formal.get(language_code, formal["de"])

    # No gender hint — fall back to neutral "Guten Tag" form (existing behavior)
    greeting = {
        "de": f"Guten Tag {contact_name},",
        "en": f"Dear {contact_name},",
        "it": f"Gentile {contact_name},",
    }
    return greeting.get(language_code, greeting["de"])


def _strip_trailing_signoff(text: str) -> str:
    """Remove trailing sign-off and candidate name from body text.

    The LLM sometimes includes sign-off despite prompt instructions not to.
    Since generate_pdf/generate_docx append their own sign-off, duplicates occur.
    """
    signoff_phrases = [
        "freundliche grüsse", "freundliche grüsse,", "freundliche gruesse",
        "kind regards", "kind regards,", "best regards", "best regards,",
        "cordiali saluti", "cordiali saluti,",
        "mit freundlichen grüssen", "mit freundlichen grüssen,",
    ]
    candidate_name = _load_profile().personal.name.lower()

    lines = text.rstrip().split("\n")
    # Walk backwards and strip sign-off lines + candidate name + empty lines
    while lines:
        last = lines[-1].strip().lower()
        if not last or last in signoff_phrases or last == candidate_name:
            lines.pop()
        else:
            break
    return "\n".join(lines)


def generate_pdf(text: str, job: dict, language_code: str, output_path: Path, contact_name: str | None = None, subtitle: str = "Revenue Operations Specialist | CRM • SAP • Power BI", letter_date: str | None = None, gender_hint: str | None = None) -> Path:
    """Generate a professional cover letter PDF."""
    pdf = CoverLetterPDF(language_code, subtitle=subtitle)
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # Date
    pdf.set_font("Arial", "", 10)
    today = datetime.strptime(letter_date, "%Y-%m-%d").strftime("%d.%m.%Y") if letter_date else datetime.now().strftime("%d.%m.%Y")
    pdf.cell(0, 6, today, align="R", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    # Company info
    company = job.get("company", "")
    location = job.get("location", "")
    if company:
        pdf.set_font("Arial", "B", 10)
        pdf.cell(0, 6, company, new_x="LMARGIN", new_y="NEXT")
    if location:
        pdf.set_font("Arial", "", 10)
        pdf.cell(0, 6, location, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    # Subject line (cleaned title — no URLs, percentages, or location suffixes)
    clean_title = _clean_title_for_subject(job.get("title", ""))
    subject_map = {
        "de": f"Bewerbung als {clean_title}",
        "en": f"Application for {clean_title}",
        "it": f"Candidatura per {clean_title}",
    }
    pdf.set_font("Arial", "B", 11)
    pdf.cell(0, 8, subject_map.get(language_code, subject_map["de"]), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    # Greeting (personalized if contact name found)
    greeting = _greeting_for_contact(contact_name, language_code, gender_hint=gender_hint)
    pdf.set_font("Arial", "", 10)
    pdf.cell(0, 6, greeting, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)

    # Body text (strip any LLM-generated sign-off to avoid duplicates)
    pdf.set_font("Arial", "", 10)
    text = _strip_trailing_signoff(text)
    paragraphs = text.strip().split("\n\n")
    for paragraph in paragraphs:
        # Clean up any single newlines within a paragraph
        paragraph = paragraph.replace("\n", " ").strip()
        if paragraph:
            pdf.multi_cell(0, 5.5, paragraph)
            pdf.ln(3)

    # Sign-off
    pdf.ln(3)
    signoff_map = {
        "de": "Freundliche Grüsse",
        "en": "Kind regards",
        "it": "Cordiali saluti",
    }
    pdf.cell(0, 6, signoff_map.get(language_code, signoff_map["de"]), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(6)
    pdf.set_font("Arial", "B", 10)
    pdf.cell(0, 6, _load_profile().personal.name, new_x="LMARGIN", new_y="NEXT")

    # Save PDF
    output_path.parent.mkdir(parents=True, exist_ok=True)
    pdf.output(str(output_path))
    return output_path


def generate_docx(text: str, job: dict, language_code: str, output_path: Path, contact_name: str | None = None, subtitle: str = "Revenue Operations Specialist | CRM • SAP • Power BI", letter_date: str | None = None, gender_hint: str | None = None) -> Path:
    """Generate a professional cover letter DOCX (editable Word document)."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page margins (2cm all around, like PDF)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Header: Name
    p = doc.add_paragraph()
    _p = _load_profile().personal
    run = p.add_run(_p.name)
    run.bold = True
    run.font.size = Pt(14)
    p.space_after = Pt(0)

    # Subtitle
    p = doc.add_paragraph()
    run = p.add_run(subtitle)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x64, 0x64, 0x64)
    p.space_after = Pt(0)

    # Contact info
    p = doc.add_paragraph()
    run = p.add_run(f"{_p.email} | {_p.phone} | {_p.location}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x64, 0x64, 0x64)
    p.space_after = Pt(2)

    # Separator line
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    from docx.oxml.ns import qn
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '4',
        qn('w:space'): '1',
        qn('w:color'): 'C8C8C8',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Date (right-aligned)
    today = datetime.strptime(letter_date, "%Y-%m-%d").strftime("%d.%m.%Y") if letter_date else datetime.now().strftime("%d.%m.%Y")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(today)
    run.font.size = Pt(10)
    p.space_after = Pt(4)

    # Company info
    company = job.get("company", "")
    location = job.get("location", "")
    if company:
        p = doc.add_paragraph()
        run = p.add_run(company)
        run.bold = True
        run.font.size = Pt(10)
        p.space_after = Pt(0)
    if location:
        p = doc.add_paragraph()
        run = p.add_run(location)
        run.font.size = Pt(10)
        p.space_after = Pt(4)

    # Subject line (cleaned title — no URLs, percentages, or location suffixes)
    clean_title = _clean_title_for_subject(job.get("title", ""))
    subject_map = {
        "de": f"Bewerbung als {clean_title}",
        "en": f"Application for {clean_title}",
        "it": f"Candidatura per {clean_title}",
    }
    p = doc.add_paragraph()
    run = p.add_run(subject_map.get(language_code, subject_map["de"]))
    run.bold = True
    run.font.size = Pt(11)
    p.space_after = Pt(8)

    # Greeting
    greeting = _greeting_for_contact(contact_name, language_code, gender_hint=gender_hint)
    p = doc.add_paragraph()
    run = p.add_run(greeting)
    run.font.size = Pt(10)
    p.space_after = Pt(4)

    # Body paragraphs (strip any LLM-generated sign-off to avoid duplicates)
    text = _strip_trailing_signoff(text)
    paragraphs = text.strip().split("\n\n")
    for paragraph in paragraphs:
        paragraph = paragraph.replace("\n", " ").strip()
        if paragraph:
            p = doc.add_paragraph()
            run = p.add_run(paragraph)
            run.font.size = Pt(10)
            p.space_after = Pt(6)

    # Sign-off
    signoff_map = {
        "de": "Freundliche Grüsse",
        "en": "Kind regards",
        "it": "Cordiali saluti",
    }
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    run = p.add_run(signoff_map.get(language_code, signoff_map["de"]))
    run.font.size = Pt(10)
    p.space_after = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run(_load_profile().personal.name)
    run.bold = True
    run.font.size = Pt(10)

    # Save DOCX
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


# ---------------------------------------------------------------------------
# Post-generation quality checks
# ---------------------------------------------------------------------------

# Numbers that are allowed in generated text (loaded from ABOUTME.md)
ALLOWED_NUMBERS = _load_profile().allowed_numbers_set()
# Date/year numbers that are always safe
SAFE_YEAR_PATTERN = re.compile(r'^20[12]\d$')  # 2016-2029
# Percentages: only these base numbers may appear with "%" (from profile metrics)
# e.g. "20" is allowed as "20%" because profile has "20% data integrity improvement"
_ALLOWED_PERCENTAGES = {"20"}
# Product version numbers — always safe regardless of candidate profile
# (e.g. "Dynamics 365", "Office 365", "SAP 2021", "Power Platform 2024")
_SAFE_PRODUCT_NUMBERS = {"365", "360", "2016", "2019", "2021", "2023", "2024", "2025"}

AI_GIVEAWAY_WORDS = {
    "en": [
        "delve", "utilize", "leverage", "moreover", "furthermore",
        "compelling", "robust", "streamline", "facilitate", "paramount",
        "embark", "navigate", "multifaceted", "transformative", "foster",
        "elevate", "realm", "bolster", "forge", "unleash", "captivate",
        "results-driven", "operational excellence", "proven expertise",
        "proven track record", "in short",
    ],
    "de": [
        "darüber hinaus", "zweifellos", "massgeblich", "wegweisend",
        "richtungsweisend", "bahnbrechend", "ganzheitlich",
        "ergebnisorientiert",
    ],
}


def _validate_no_hallucinated_metrics(text: str, extra_context: str = "") -> list[str]:
    """Find numbers in generated text that aren't in the candidate profile or company context."""
    # Extract numbers from company context (these are valid to mention)
    context_numbers = set()
    context_percentages = set()  # Track percentages separately — "30 employees" ≠ "30%"
    # Regex: match numbers with optional %, using lookahead instead of trailing \b
    # (trailing \b fails after % since % is a non-word character)
    _NUM_RE = r'\b\d[\d,\'\.]*\+?%?'
    if extra_context:
        for num in re.findall(_NUM_RE, extra_context):
            clean_ctx = num.rstrip('%+').replace(',', '').replace("'", '').replace('.', '')
            if '%' in num:
                context_percentages.add(clean_ctx)
                context_percentages.add(num.rstrip('%+'))
                # Also add to context_numbers — handles format mismatches where the LLM
                # echoes a company context percentage but strips/adds formatting variants
                context_numbers.add(clean_ctx)
            else:
                context_numbers.add(clean_ctx)
                context_numbers.add(num.rstrip('%+'))

    # Extract all numbers/percentages from generated text
    generated = set(re.findall(_NUM_RE, text))
    violations = []
    for num in generated:
        clean = num.rstrip('%+').replace(',', '').replace("'", '').replace('.', '')
        is_percentage = '%' in num

        # Percentages get stricter validation — only allowed if percentage appears in context
        if is_percentage:
            if clean in _ALLOWED_PERCENTAGES:
                continue
            if clean in context_percentages or num.rstrip('%+') in context_percentages:
                continue
            violations.append(num)
            continue

        # Non-percentage numbers: check against full allowed set
        if clean in ALLOWED_NUMBERS:
            continue
        if SAFE_YEAR_PATTERN.match(clean):
            continue
        if num.rstrip('%+') in ALLOWED_NUMBERS:
            continue
        if clean in context_numbers or num.rstrip('%+') in context_numbers:
            continue
        if clean in _SAFE_PRODUCT_NUMBERS:
            continue
        violations.append(num)
    return violations


def _filter_meta_commentary(text: str) -> tuple[str, list[str]]:
    """Remove meta-writing commentary that LLMs insert to manipulate burstiness scores.

    Examples: "Short sentences work.", "Clarity matters.", "Less is more.",
    "Results speak.", "Numbers tell the story."
    These are writing-advice artifacts, not cover letter content.
    """
    # Match very short standalone sentences (2-5 words) that read as style commentary
    # rather than substantive claims. Only remove if they don't contain numbers or
    # domain-specific terms (those are likely real content).
    meta_patterns = [
        r'(?<=[.!?])\s+(?:Short sentences work|Clarity matters|Less is more|Results speak|'
        r'Numbers tell the story|That matters|This counts|Details matter|Precision counts|'
        r'Quality matters|Impact matters|Data speaks|Facts matter|Actions count|'
        r'Kurze Sätze wirken|Klarheit zählt|Das zählt|Qualität zählt|Fakten sprechen'
        r')\.?',
    ]
    found = []
    for pat in meta_patterns:
        matches = re.findall(pat, text, re.IGNORECASE)
        if matches:
            found.extend(m.strip() for m in matches)
            text = re.sub(pat, '', text, flags=re.IGNORECASE)
    text = re.sub(r'  +', ' ', text).strip()
    return text, found


def _filter_ai_giveaway_words(text: str, lang_code: str) -> tuple[str, list[str]]:
    """Remove AI giveaway words and return (cleaned_text, found_words)."""
    found = []
    words = AI_GIVEAWAY_WORDS.get(lang_code, []) + AI_GIVEAWAY_WORDS.get("en", [])
    for word in words:
        pattern = re.compile(re.escape(word), re.IGNORECASE)
        if pattern.search(text):
            found.append(word)
            # Remove the word, cleaning up extra spaces
            text = pattern.sub('', text)
    # Clean up double spaces
    text = re.sub(r'  +', ' ', text).strip()
    return text, found


def _fix_grammar_after_removal(text: str, openrouter_key: str | None, gemini_key: str | None) -> str:
    """[H5] LLM pass to fix grammar errors in cover letter text (~100 tokens).

    Called on every cover letter to catch both AI-word-removal artifacts and
    independent grammar errors from the LLM.
    """
    prompt = (
        "Fix any grammatical errors in the following text. Do not change meaning, "
        "style, tone, or length. Do not add new content. Return ONLY the corrected text, "
        "nothing else.\n\n" + text
    )
    try:
        fixed, _ = _call_llm(openrouter_key, gemini_key, prompt, temperature=0.1)
        # Sanity: fixed text should be similar length (within 20%)
        orig_wc = len(text.split())
        fixed_wc = len(fixed.split())
        if fixed_wc < orig_wc * 0.8 or fixed_wc > orig_wc * 1.2:
            log.warning(f"  Grammar fix changed length too much ({orig_wc} → {fixed_wc} words), keeping original")
            return text
        return fixed
    except Exception as e:
        log.warning(f"  Grammar fix LLM call failed: {e}, keeping original")
        return text


def _enforce_swiss_german(text: str) -> str:
    """Replace German ß with Swiss ss."""
    return text.replace('ß', 'ss')


def _validate_word_count(text: str, min_words: int = 200, max_words: int = 260) -> int:
    """Check that the cover letter body is within the ideal word count range.
    Returns the word count. Logs a warning if out of range."""
    words = text.split()
    count = len(words)
    if count < min_words:
        log.warning(f"  Cover letter too short: {count} words (target: {min_words}-{max_words})")
    elif count > max_words:
        log.warning(f"  Cover letter too long: {count} words (target: {min_words}-{max_words})")
    else:
        log.info(f"  Word count: {count} (target: {min_words}-{max_words})")
    return count


def _check_sentence_burstiness(text: str) -> float:
    """Measure sentence length variation — higher StdDev = more human-like.
    Returns standard deviation of sentence word counts.
    AI-generated text typically has StdDev < 5; human text > 8."""
    import statistics
    sentences = re.split(r'(?<=[.!?])\s+', text.strip())
    lengths = [len(s.split()) for s in sentences if len(s.split()) > 0]
    if len(lengths) < 3:
        return 0.0
    sd = statistics.stdev(lengths)
    if sd < 5:
        log.warning(f"  Low sentence burstiness (StdDev={sd:.1f}) — text may seem AI-generated. Target: >8")
    else:
        log.info(f"  Sentence burstiness OK (StdDev={sd:.1f})")
    return sd


def _validate_keyword_coverage(text: str, key_matches: list[str]) -> tuple[float, list[str]]:
    """Check that key job terms from evaluation appear in the generated text.

    Uses fuzzy matching: a compound key_match like "data analysis Power BI SQL"
    counts as matched if any significant word from it appears in the text.
    Returns (coverage_pct, missing_terms).
    """
    if not key_matches:
        return 1.0, []
    text_lower = text.lower()
    # Trivial words to skip when splitting compound key_matches
    skip_words = {"and", "or", "the", "for", "with", "in", "of", "a", "an", "to", "on"}
    missing = []
    for km in key_matches:
        km_lower = km.lower()
        # Check exact match first
        if km_lower in text_lower:
            continue
        # Fuzzy: check if any significant word from the key_match appears
        words = [w for w in km_lower.split() if w not in skip_words and len(w) > 2]
        if words and any(w in text_lower for w in words):
            continue
        missing.append(km)
    coverage = 1.0 - (len(missing) / len(key_matches)) if key_matches else 1.0
    if coverage < 0.5:
        log.warning(f"  Low keyword coverage: {coverage:.0%} — missing: {missing}")
    else:
        log.info(f"  Keyword coverage: {coverage:.0%} ({len(key_matches) - len(missing)}/{len(key_matches)})")
    return coverage, missing


_FOREIGN_SCRIPT_RE = re.compile(
    r'[\u4e00-\u9fff'   # CJK Unified Ideographs (Chinese/Japanese/Korean)
    r'\u3040-\u30ff'    # Hiragana + Katakana
    r'\u0600-\u06ff'    # Arabic
    r'\u0590-\u05ff'    # Hebrew
    r'\u0900-\u097f'    # Devanagari
    r'\uff00-\uffef]'   # Fullwidth / Halfwidth forms
)


def _validate_no_foreign_script(text: str) -> bool:
    """Return True if clean, False if non-Latin script detected (corruption signal)."""
    match = _FOREIGN_SCRIPT_RE.search(text)
    if match:
        # Log the surrounding context for debugging
        start = max(0, match.start() - 20)
        end = min(len(text), match.end() + 20)
        log.warning(f"  Non-Latin characters detected: ...{text[start:end]}...")
    return not bool(match)


def _validate_company_mention(text: str, company: str) -> bool:
    """Check that the company is mentioned in the cover letter body."""
    if not company:
        return True
    # Strip legal suffixes and check for the core company name
    clean = strip_legal_suffixes(company)
    # Also strip TLDs (e.g. "Huzzle.com" → "Huzzle")
    clean = re.sub(r'\.(com|ch|de|net|org|io|ai|co)\b', '', clean, flags=re.IGNORECASE).strip()
    if not clean or len(clean) < 3:
        return True  # Can't validate very short names
    text_lower = text.lower()
    if clean.lower() in text_lower:
        return True
    # Fallback: check if any significant word from the company name appears in text
    # (handles long names like "Stadler Signalling Deutschland" → matches on "Stadler")
    stop_words = {'the', 'and', 'for', 'of', 'des', 'der', 'die', 'und', 'für', 'von'}
    words = [w for w in clean.split() if len(w) > 2 and w.lower() not in stop_words]
    if words and any(w.lower() in text_lower for w in words):
        return True
    log.warning(f"  Company name '{clean}' not found in cover letter text — may be generic")
    return False


# ---------------------------------------------------------------------------
# Checkpointing (resume on interruption)
# ---------------------------------------------------------------------------

def _get_job_id(job: dict) -> str:
    """Create a unique ID for a job based on title + company + URL."""
    return generate_job_id(job.get("title", ""), job.get("company", ""), job.get("url", ""))


def _load_checkpoint() -> dict:
    """Load checkpoint tracking which jobs have been processed."""
    if CHECKPOINT_FILE.exists():
        try:
            with open(CHECKPOINT_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except (json.JSONDecodeError, OSError):
            return {"processed": {}}
    return {"processed": {}}


def _save_checkpoint(checkpoint: dict):
    """Persist checkpoint to disk."""
    CHECKPOINT_FILE.parent.mkdir(parents=True, exist_ok=True)
    with open(CHECKPOINT_FILE, "w", encoding="utf-8") as f:
        json.dump(checkpoint, f, ensure_ascii=False, indent=2)


# ---------------------------------------------------------------------------
# Main pipeline
# ---------------------------------------------------------------------------

_AGENCY_KEYWORDS = [
    "personalberatung", "recruitment", "recruiting", "staffing", "headhunt",
    "personal ag", "personal gmbh", "personalvermittlung", "randstad",
    "adecco", "manpower", "hays", "robert half", "michael page", "kienbaum",
    "egon zehnder", "huzzle",
]


def _detect_actual_employer(company: str, description: str) -> str | None:
    """If company is a recruitment agency, extract the actual employer from the description."""
    if not company or not description:
        return None
    company_lower = company.lower()
    is_agency = any(kw in company_lower for kw in _AGENCY_KEYWORDS)
    if not is_agency:
        return None

    # Pattern 1: Bold employer name at start of description (LinkedIn markdown: **COMPANY**)
    m = re.match(r'\s*\*\*([A-ZÄÖÜ][^*]{2,40})\*\*', description)
    if m:
        name = m.group(1).strip()
        # Filter out labels like "Job Type:", "Location:", etc.
        if not name.endswith(":") and not re.match(
            r'^(Job|Location|Engagement|About|Your|Our|The|Key|Main|Role|Tasks|Requirements|What|How|Why|Overview|Position|Company|Summary)\b',
            name,
        ):
            return name

    # Pattern 2: "Für unseren Kunden [COMPANY]" / "For our client [COMPANY]"
    m = re.search(
        r'(?:für\s+unsere[n]?\s+Kund(?:en|in)|for\s+our\s+client)\s+([A-ZÄÖÜ][A-Za-zÄÖÜäöüß&. -]{2,40})',
        description, re.IGNORECASE,
    )
    if m:
        return m.group(1).strip().rstrip(".,;:")

    # Pattern 3: "Bei [COMPANY]" / "At [COMPANY]" in first 300 chars
    first_block = description[:300]
    m = re.search(r'(?:^|\n)\s*(?:Bei|At)\s+([A-ZÄÖÜ][A-Za-zÄÖÜäöüß&. -]{2,40})', first_block)
    if m:
        name = m.group(1).strip().rstrip(".,;:")
        if name.lower() != company_lower:
            return name

    return None


def generate_cover_letter_for_job(
    openrouter_key: str | None,
    gemini_key: str | None,
    job: dict,
    letter_date: str | None = None,
) -> dict:
    """Generate a cover letter for a single job. Returns result dict."""
    title = job.get("title", "Unknown")
    company = job.get("company", "Unknown")
    location = job.get("location", "Unknown")
    description = job.get("description", "")
    key_matches = job.get("key_matches", [])
    key_gaps = job.get("key_gaps", [])

    # Detect recruitment agency — extract actual employer from description
    actual_employer = _detect_actual_employer(company, description)
    # Compare base names (strip .com, AG, GmbH, etc.)
    _strip = lambda s: re.sub(r'\.(com|ch|de|net|org)$', '', s.lower().strip()).strip()
    if actual_employer and _strip(actual_employer) != _strip(company):
        log.info(f"  Recruitment agency detected: '{company}' → actual employer: '{actual_employer}'")
        company = actual_employer  # Use actual employer for letter and research

    # Detect language from title + description
    lang_code, lang_name = detect_language(title, description)
    log.info(f"  Language detected: {lang_name}")

    # Try to find a contact person name: regex first (free), LLM fallback (cheap)
    contact_name = _extract_contact(description, openrouter_key=openrouter_key, gemini_key=gemini_key)
    # Best-effort gender hint for the salutation — purely additive: if missing, the
    # greeting falls back to today's neutral "Guten Tag {name}," form.
    gender_hint = detect_honorific(description) if contact_name else None
    if gender_hint:
        log.info(f"  Honorific detected: {gender_hint}")

    # Research company for personalization (cached, cheap)
    job_url = job.get("url", "")
    company_context = _research_company(company, description, openrouter_key, gemini_key, job_url=job_url, job_title=title)

    # Truncate description for prompt
    desc_for_prompt = description
    if len(desc_for_prompt) > 2000:
        desc_for_prompt = desc_for_prompt[:2000] + "..."

    # Select profile based on provider (anonymized for OpenRouter/Deepseek)
    profile = _get_profile_for_provider(openrouter_key)

    # Build prompt
    prompt = COVER_LETTER_PROMPT.format(
        profile=profile,
        title=title,
        company=company,
        location=location,
        key_matches=", ".join(key_matches) if key_matches else "General fit",
        key_gaps=", ".join(key_gaps) if key_gaps else "None identified",
        company_context=company_context if company_context else "No specific company information available — use details from the job description.",
        description=desc_for_prompt,
        language_name=lang_name,
        language_code=lang_code,
    )

    # Inject contextually relevant project hint (soft — LLM includes only if it fits naturally)
    profile_obj = _load_profile()
    relevant_projs = profile_obj.relevant_projects(key_matches, lang_code, top_n=1)
    if relevant_projs:
        p = relevant_projs[0]
        if lang_code == "de":
            proj_instruction = (
                f"Mein Projekt '{p['name']}' ({p['url']}) zeigt, wie ich [relevante Fähigkeit] in der Praxis anwende."
            )
            project_hint = (
                f"\n\nRELEVANT PROJECT (nur erwähnen, wenn es natürlich in den Fliesstext passt — NICHT als separaten Absatz!):\n"
                f"Name: {p['name']}\nTechnologien: {p['tech']}\nURL: {p['url']}\n"
                f"Beschreibung: {p['description']}\n"
                f"Beispielsatz: \"{proj_instruction}\"\n"
                f"Nur einflechten, wenn es einen konkreten Anspruch stärkt. "
                f"Kürze dabei keine anderen Absätze, um die Projekterwähnung unterzubringen."
            )
        elif lang_code == "it":
            project_hint = (
                f"\n\nPROGETTO RILEVANTE (menzionare solo se si integra naturalmente — NON come paragrafo separato!):\n"
                f"Nome: {p['name']}\nTecnologie: {p['tech']}\nURL: {p['url']}\n"
                f"Descrizione: {p['description']}\n"
                f"Solo se rafforza un'affermazione specifica. "
                f"Non accorciare gli altri paragrafi per inserire la menzione del progetto."
            )
        else:  # en
            project_hint = (
                f"\n\nRELEVANT PROJECT (mention only if it fits naturally — do NOT add as a separate paragraph):\n"
                f"Name: {p['name']}\nTech: {p['tech']}\nURL: {p['url']}\n"
                f"Description: {p['description']}\n"
                f"Example: \"My project '{p['name']}' ({p['url']}) demonstrates [specific skill] in practice.\"\n"
                f"Only include if it strengthens a specific claim. "
                f"Do not shorten other paragraphs to accommodate the project mention."
            )
        prompt += project_hint
        log.info(f"  Project hint injected: {p['name']} ({p['url']})")

    # --------------- LLM generation with quality gate retry loop ---------------
    # Try up to 3 times, bumping temperature on failure to get varied output.
    # Track best attempt (fewest quality issues) as fallback.
    MAX_QUALITY_ATTEMPTS = 3
    best_text = None
    best_subtitle = enforce_revops_subtitle(None)
    best_quality = None  # dict tracking quality metrics
    best_score = -1  # higher = better
    retry_extra = ""  # extra instruction injected on low-word-count retries

    for quality_attempt in range(MAX_QUALITY_ATTEMPTS):
        temperature = 0.3 + (quality_attempt * 0.1)  # 0.3, 0.4, 0.5

        if quality_attempt > 0:
            log.warning(f"  Quality retry {quality_attempt + 1}/{MAX_QUALITY_ATTEMPTS} (temperature={temperature:.1f})")

        # Call LLM (retry_extra appended when previous attempt was too short)
        raw_response, provider = _call_llm(openrouter_key, gemini_key, prompt + retry_extra, temperature=temperature)
        if quality_attempt == 0:
            log.info(f"  Provider: {provider} ({'anonymized' if provider == 'openrouter' else 'full profile'})")

        # Parse JSON response (with fallback for plain text)
        subtitle = enforce_revops_subtitle(None)
        cleaned = re.sub(r'^```(?:json)?\s*', '', raw_response.strip(), flags=re.MULTILINE)
        cleaned = re.sub(r'```\s*$', '', cleaned.strip(), flags=re.MULTILINE)
        text = cleaned
        try:
            result = json.loads(cleaned)
            text = result.get("text", cleaned)
            raw_subtitle = result.get("subtitle")
            subtitle = enforce_revops_subtitle(raw_subtitle)
            if raw_subtitle and raw_subtitle != subtitle:
                log.info(f"  Subtitle normalized: {raw_subtitle!r} → {subtitle!r}")
            else:
                log.info(f"  Subtitle: {subtitle}")
        except (json.JSONDecodeError, KeyError):
            log.warning("  LLM returned plain text instead of JSON, using as-is")

        # Clean up text (remove accidental markdown or headers)
        text = re.sub(r'^#+\s.*$', '', text, flags=re.MULTILINE).strip()
        text = re.sub(r'^\*\*.*\*\*$', '', text, flags=re.MULTILINE).strip()

        # --- Quality checks ---
        text = _enforce_swiss_german(text)

        violations = _validate_no_hallucinated_metrics(text, extra_context=company_context or "")
        text, ai_words_found = _filter_ai_giveaway_words(text, lang_code)
        if ai_words_found:
            log.info(f"  Removed AI giveaway words: {ai_words_found}")
        text, meta_found = _filter_meta_commentary(text)
        if meta_found:
            log.info(f"  Removed meta-commentary: {meta_found}")
        text = _fix_grammar_after_removal(text, openrouter_key, gemini_key)

        word_count = _validate_word_count(text)
        burstiness = _check_sentence_burstiness(text)
        keyword_coverage, missing_keywords = _validate_keyword_coverage(text, key_matches)
        company_mentioned = _validate_company_mention(text, company)
        script_clean = _validate_no_foreign_script(text)

        # Track quality metrics
        quality = {
            "hallucinations": violations,
            "script_clean": script_clean,
            "word_count": word_count,
            "burstiness": burstiness,
            "keyword_coverage": keyword_coverage,
            "missing_keywords": missing_keywords,
            "company_mentioned": company_mentioned,
            "ai_words_removed": ai_words_found,
        }

        # Score: higher = better (max 5). Critical = 0 pts, soft = partial credit.
        score = 0
        score += 3 if (not violations and script_clean) else 0  # hallucination/corruption = most critical
        if 200 <= word_count <= 260:
            score += 1                                   # ideal word count range
        elif 190 <= word_count <= 300:
            score += 0.5                                 # within hard bounds but not ideal
        if burstiness >= 5:
            score += 1                                   # good sentence variety
        elif burstiness >= 3:
            score += 0.5                                 # acceptable but flagged

        if score > best_score:
            best_text = text
            best_subtitle = subtitle
            best_quality = quality
            best_score = score

        # Retry only on truly critical failures worth the LLM cost:
        # - Hallucinated metrics (instant rejection risk)
        # - Severely off word count (< 190 or > 300)
        # - Very low burstiness (< 3, obviously AI)
        # - Company name not mentioned (Swiss recruiters expect it)
        critical_pass = not violations and script_clean and 190 <= word_count <= 300 and burstiness >= 3 and company_mentioned
        if critical_pass:
            break

        if quality_attempt < MAX_QUALITY_ATTEMPTS - 1:
            issues = []
            if violations:
                issues.append(f"hallucinations: {violations}")
            if not script_clean:
                issues.append("non-Latin script detected")
            if word_count < 190 or word_count > 300:
                issues.append(f"word count: {word_count}")
            if burstiness < 3:
                issues.append(f"burstiness: {burstiness:.1f}")
            if not company_mentioned:
                issues.append(f"company '{company}' not mentioned")
            log.warning(f"  Critical quality issues: {', '.join(issues)} — regenerating...")

            # Retry injections: targeted corrections for each failure mode
            if not script_clean:
                retry_extra += (
                    "\n\nCRITICAL: Your previous response contained non-Latin characters (e.g. Chinese, Arabic). "
                    "Write ONLY using standard Latin characters. No special scripts."
                )
            if not company_mentioned:
                retry_extra += (
                    f"\n\nCRITICAL: You MUST mention '{company}' by name in the letter body — this is mandatory. "
                    f"Weave it naturally into your opening or body paragraph (e.g. 'Bei {company} sehe ich...' "
                    f"or 'Die Position bei {company} bietet...'). Do not omit the company name."
                )
                log.warning(f"  Adding company name injection for next attempt: {company}")

            # Floor injection: if word count too low, remind LLM on next attempt
            if word_count < 210:
                retry_extra = (
                    f"\n\nCRITICAL CORRECTION: Your previous attempt was only {word_count} words — far too short. "
                    f"You MUST write at least 260 words. Expand each body paragraph with one additional specific detail "
                    f"(HOW you did it, what tool you used, what the result was). Do not skip this."
                )
                log.warning(f"  Adding word count floor injection for next attempt ({word_count} words was too short)")

            time.sleep(1.0)  # brief pause before retry

    # Use the best attempt
    text = best_text
    subtitle = best_subtitle
    quality = best_quality

    if best_score < 5:
        log.warning(f"  QUALITY WARNING: Best attempt scored {best_score}/5 after {quality_attempt + 1} tries — review before sending!")

    # Log whether the project link made it into the final text (soft check)
    if relevant_projs:
        proj_url = relevant_projs[0]["url"]
        if proj_url in (best_text or ""):
            log.info(f"  [OK] Project link included in cover letter: {proj_url}")
        else:
            log.info(f"  [INFO] Project link not included (LLM chose not to): {proj_url}")

    # Generate PDF
    job_id = _get_job_id(job)
    safe_company = sanitize_filename(company) if company else "Unknown"
    safe_title = sanitize_filename(clean_job_title(title)) if title else "Unknown"
    folder_name = f"{job_id}_{safe_company}_{safe_title}"
    safe_name = sanitize_filename(_load_profile().personal.name)
    output_path = OUTPUT_DIR / folder_name / f"Cover_Letter_{safe_name}.pdf"

    pdf_path = generate_pdf(text, job, lang_code, output_path, contact_name=contact_name, subtitle=subtitle, letter_date=letter_date, gender_hint=gender_hint)
    log.info(f"  PDF saved: {pdf_path}")

    # Generate DOCX (editable version)
    docx_output = output_path.with_suffix(".docx")
    docx_path = generate_docx(text, job, lang_code, docx_output, contact_name=contact_name, subtitle=subtitle, letter_date=letter_date, gender_hint=gender_hint)
    log.info(f"  DOCX saved: {docx_path}")

    return {
        "title": title,
        "company": company,
        "language": lang_name,
        "lang_code": lang_code,
        "subtitle": subtitle,
        "pdf_path": str(pdf_path),
        "docx_path": str(docx_path),
        "text_length": len(text),
        "quality": {
            "score": best_score,
            "max_score": 5,  # 3 (hallucinations) + 1 (word count) + 1 (burstiness)
            "hallucinations": quality.get("hallucinations", []),
            "word_count": quality.get("word_count", 0),
            "burstiness": round(quality.get("burstiness", 0), 1),
            "keyword_coverage": round(quality.get("keyword_coverage", 0), 2),
            "missing_keywords": quality.get("missing_keywords", []),
            "company_mentioned": quality.get("company_mentioned", True),
        },
    }


def _load_ready_to_apply_from_sheet(sheet_name: str = "Swiss Job Search Pipeline") -> list[dict]:
    """Read rows with Status=Ready_to_Apply from the Google Sheet.

    Returns job dicts compatible with generate_cover_letter_for_job().
    Tries to enrich description from scored_jobs.json on volume (full text);
    falls back to the truncated sheet description if not found.
    """
    from execution.write_jobs_to_sheet import authenticate, _sheets_api_call

    client = authenticate()
    try:
        ws = _sheets_api_call(client.open, sheet_name).sheet1
    except Exception as e:
        log.error(f"[sheet-triggered] Cannot open sheet '{sheet_name}': {e}")
        return []

    rows = _sheets_api_call(ws.get_all_values)
    if not rows or len(rows) < 2:
        return []

    header = [h.strip() for h in rows[0]]

    def col(name: str) -> int | None:
        try:
            return header.index(name)
        except ValueError:
            return None

    ci = {name: col(name) for name in (
        "Job_ID", "Title", "Company", "Location", "Source",
        "Score", "Key Matches", "Key Gaps", "Description", "URL",
        "Date Posted", "Status", "Reasoning",
    )}

    # Load scored_jobs.json for full descriptions (best-effort)
    scored_by_id: dict[str, dict] = {}
    if INPUT_FILE.exists():
        try:
            with open(INPUT_FILE, "r", encoding="utf-8") as f:
                for j in json.load(f):
                    scored_by_id[_get_job_id(j)] = j
        except Exception:
            pass

    def cell(row: list, name: str) -> str:
        idx = ci.get(name)
        return row[idx].strip() if idx is not None and idx < len(row) else ""

    jobs = []
    for row in rows[1:]:
        status = cell(row, "Status").lower()
        if status != "ready_to_apply":
            continue

        job_id = cell(row, "Job_ID")
        url    = cell(row, "URL")
        title  = cell(row, "Title")
        company = cell(row, "Company")

        # Prefer full description from scored_jobs.json
        if job_id and job_id in scored_by_id:
            base = scored_by_id[job_id]
            description = base.get("description", cell(row, "Description"))
        else:
            description = cell(row, "Description")

        try:
            score = int(cell(row, "Score"))
        except ValueError:
            score = 0

        jobs.append({
            "job_id": job_id,
            "title": title,
            "company": company,
            "location": cell(row, "Location"),
            "source": cell(row, "Source"),
            "url": url,
            "description": description,
            "score": score,
            "key_matches": [k.strip() for k in cell(row, "Key Matches").split(",") if k.strip()],
            "key_gaps": [k.strip() for k in cell(row, "Key Gaps").split(",") if k.strip()],
            "reasoning": cell(row, "Reasoning"),
            "date_posted": cell(row, "Date Posted"),
        })

    log.info(f"[sheet-triggered] Found {len(jobs)} Ready_to_Apply rows in '{sheet_name}'")
    return jobs


def main():
    parser = argparse.ArgumentParser(description="Generate cover letters for scored jobs")
    parser.add_argument("--limit", type=int, default=0, help="Limit number of cover letters (0 = all)")
    parser.add_argument("--min-score", type=int, default=DEFAULT_MIN_SCORE, help="Minimum score threshold (default: 6)")
    parser.add_argument("--reset-checkpoint", action="store_true", help="Clear checkpoint and re-generate all")
    parser.add_argument("--letter-date", type=str, default=None, help="Override letter date (YYYY-MM-DD format, e.g. 2026-04-01)")
    parser.add_argument("--job-id", type=str, default=None, help="Generate only for a specific job_id (e.g. J-ab2e15)")
    parser.add_argument("--sheet-triggered", action="store_true",
                        help="Read Ready_to_Apply rows from Sheet instead of scored_jobs.json. "
                             "Updates Status → APPLYING after generation.")
    args = parser.parse_args()

    openrouter_key = os.getenv("OPEN_ROUTER_API_KEY")
    gemini_key = os.getenv("GOOGLE_AI_STUDIO_API_KEY")

    if not openrouter_key and not gemini_key:
        log.error("No API keys found. Set OPEN_ROUTER_API_KEY or GOOGLE_AI_STUDIO_API_KEY in .env")
        sys.exit(1)

    # Load or reset checkpoint
    if args.reset_checkpoint and CHECKPOINT_FILE.exists():
        CHECKPOINT_FILE.unlink()
        log.info("Checkpoint cleared — will re-generate all cover letters")

    checkpoint = _load_checkpoint()

    # --sheet-triggered: read Ready_to_Apply rows directly from Sheet
    if args.sheet_triggered:
        scored_jobs = _load_ready_to_apply_from_sheet()
        if not scored_jobs:
            log.info("[sheet-triggered] No Ready_to_Apply rows. Nothing to do.")
            return
        log.info(f"[sheet-triggered] Processing {len(scored_jobs)} Ready_to_Apply job(s)")
        eligible_jobs = scored_jobs
    else:
        if not INPUT_FILE.exists():
            log.error(f"Input file not found: {INPUT_FILE}")
            log.error("Run execution/evaluate_jobs.py first (Stage 2)")
            sys.exit(1)

        with open(INPUT_FILE, "r", encoding="utf-8") as f:
            scored_jobs = json.load(f)

        log.info(f"Loaded {len(scored_jobs)} scored jobs from {INPUT_FILE}")

        # Filter by minimum score
        eligible_jobs = [j for j in scored_jobs if j.get("score", 0) >= args.min_score]
        log.info(f"Jobs scoring >= {args.min_score}: {len(eligible_jobs)}")

    if args.job_id:
        eligible_jobs = [j for j in eligible_jobs if _get_job_id(j) == args.job_id]
        if not eligible_jobs:
            log.error(f"Job ID {args.job_id} not found in scored_jobs.json (or below --min-score {args.min_score})")
            sys.exit(1)
        log.info(f"--job-id filter: targeting {args.job_id}")

    if args.limit > 0:
        eligible_jobs = eligible_jobs[:args.limit]
        log.info(f"Limiting to {len(eligible_jobs)} cover letters")

    if not eligible_jobs:
        log.info("No jobs meet the score threshold. Nothing to generate.")
        return

    # Filter out jobs with empty company (can't produce professional cover letters)
    no_company = [j for j in eligible_jobs if not j.get("company", "").strip()]
    if no_company:
        log.warning(f"Skipping {len(no_company)} jobs with no company name (unprofessional output)")
        for j in no_company:
            log.warning(f"  - '{j.get('title', '?')}' (score: {j.get('score', '?')})")
        eligible_jobs = [j for j in eligible_jobs if j.get("company", "").strip()]

    # Filter out already-processed jobs (checkpoint + folder-existence safety net)
    already_done = checkpoint.get("processed", {})
    pending_jobs = []
    folder_synced = 0
    for job in eligible_jobs:
        job_id = _get_job_id(job)
        if job_id in already_done:
            log.debug(f"  Skipping (checkpoint): {job.get('title', '?')} at {job.get('company', '?')}")
            continue
        existing_folder = find_application_folder(OUTPUT_DIR, job_id)
        if existing_folder and has_cover_letter_outputs(existing_folder):
            log.info(f"  Skipping (folder exists): {job.get('title', '?')} at {existing_folder.name}")
            existing_pdf = next(iter(existing_folder.glob("Cover_Letter_*.pdf")), None)
            already_done[job_id] = {
                "title": job.get("title", ""),
                "company": job.get("company", ""),
                "pdf_path": str(existing_pdf) if existing_pdf else "",
            }
            folder_synced += 1
            continue
        pending_jobs.append(job)

    if folder_synced > 0:
        checkpoint["processed"] = already_done
        _save_checkpoint(checkpoint)
        log.info(f"Synced {folder_synced} pre-existing folder(s) to checkpoint (no LLM cost)")

    skipped = len(eligible_jobs) - len(pending_jobs)
    if skipped > 0:
        log.info(f"Skipping {skipped} already-generated cover letters (checkpoint + folder dedup)")

    if not pending_jobs:
        log.info("All eligible jobs already have cover letters. Use --reset-checkpoint to re-generate.")
        return

    log.info(f"Generating cover letters for {len(pending_jobs)} jobs")

    # Generate cover letters
    results = []
    for i, job in enumerate(pending_jobs):
        log.info(f"[{i + 1}/{len(pending_jobs)}] Generating cover letter: {job.get('title', '?')} at {job.get('company', '?')} (score: {job.get('score', '?')})")

        try:
            result = generate_cover_letter_for_job(openrouter_key, gemini_key, job, letter_date=args.letter_date)
            results.append(result)

            # Save to checkpoint (includes subtitle + language for CV consistency)
            job_id = _get_job_id(job)
            checkpoint.setdefault("processed", {})[job_id] = {
                "title": job.get("title", "?"),
                "company": job.get("company", "?"),
                "pdf_path": result.get("pdf_path", ""),
                "subtitle": result.get("subtitle", ""),
                "language": result.get("lang_code", ""),
            }
            _save_checkpoint(checkpoint)

            # Update Google Sheet with CL status
            try:
                from execution.write_jobs_to_sheet import update_job_columns
                q = result.get("quality", {})
                update_job_columns(
                    sheet_name="Swiss Job Search Pipeline",
                    job_url=job.get("url", ""),
                    updates={
                        "CL_Generated": "Yes",
                        "CL_Quality_Score": f"{q.get('score', 0):.3g}/{q.get('max_score', 5)} ({q.get('word_count', '?')}w)",
                    },
                )
            except Exception as e:
                log.debug(f"  Sheet update skipped: {e}")

            # Upload application folder to Google Drive
            try:
                from execution.drive_upload import upload_application_folder
                pdf_path = result.get("pdf_path", "")
                if pdf_path:
                    drive_url = upload_application_folder(Path(pdf_path).parent)
                    if drive_url:
                        log.info(f"  Drive: {drive_url}")
            except Exception as e:
                log.debug(f"  Drive upload skipped: {e}")

            # --sheet-triggered: update Status → APPLYING so CV stage picks it up
            if args.sheet_triggered:
                try:
                    from execution.write_jobs_to_sheet import update_job_columns
                    update_job_columns(
                        sheet_name="Swiss Job Search Pipeline",
                        job_url=job.get("url", ""),
                        updates={"Status": "APPLYING"},
                    )
                    log.info(f"  Sheet: Status → APPLYING")
                except Exception as e:
                    log.debug(f"  Status update skipped: {e}")
        except Exception as e:
            log.error(f"  Failed: {e}")
            results.append({
                "title": job.get("title", "?"),
                "company": job.get("company", "?"),
                "error": str(e),
            })

        # Rate limit between calls
        if i < len(pending_jobs) - 1:
            time.sleep(2.0)

    # Summary
    successful = [r for r in results if "pdf_path" in r]
    failed = [r for r in results if "error" in r]

    log.info(f"\n{'='*60}")
    log.info(f"Cover Letter Generation Complete")
    log.info(f"  Generated: {len(successful)}/{len(results)}")
    if skipped > 0:
        log.info(f"  Skipped (checkpoint): {skipped}")
    if failed:
        log.info(f"  Failed: {len(failed)}")

    # Language breakdown
    lang_counts = {}
    for r in successful:
        lang = r.get("language", "Unknown")
        lang_counts[lang] = lang_counts.get(lang, 0) + 1
    log.info(f"  Languages: {lang_counts}")

    log.info(f"  Output directory: {OUTPUT_DIR}")
    for r in successful:
        log.info(f"  - {r['company']}: {r['title']} ({r['language']}) [{r.get('subtitle', '?')}]")

    # Quality report
    if successful:
        log.info(f"\n{'='*60}")
        log.info("QUALITY REPORT:")
        quality_report = []
        for r in successful:
            q = r.get("quality", {})
            score = q.get("score", 0)
            max_s = q.get("max_score", 5)
            issues = []
            if q.get("hallucinations"):
                issues.append(f"hallucinations: {q['hallucinations']}")
            wc = q.get("word_count", 0)
            if wc < 200 or wc > 260:
                issues.append(f"word count: {wc}")
            burst = q.get("burstiness", 0)
            if burst < 5:
                issues.append(f"burstiness: {burst}")
            if not q.get("company_mentioned", True):
                issues.append("company not mentioned")
            kc = q.get("keyword_coverage", 1.0)
            if kc < 0.5:
                issues.append(f"keyword coverage: {kc:.0%}")

            if not issues:
                status = "PASS"
                icon = "+"
            elif issues and any("hallucination" in i for i in issues):
                status = "NEEDS REVIEW"
                icon = "X"
            elif issues:
                status = "WARNING"
                icon = "!"
            else:
                status = "PASS"
                icon = "+"

            log.info(f"  [{icon}] {r['company']} - {r['title']}: {status}" +
                     (f" ({', '.join(issues)})" if issues else ""))

            quality_report.append({
                "company": r["company"],
                "title": r["title"],
                "status": status,
                "score": score,
                "issues": issues,
                "quality": q,
            })

        # Write quality report to JSON
        report_path = TMP_DIR / "quality_report.json"
        with open(report_path, "w", encoding="utf-8") as f:
            json.dump(quality_report, f, ensure_ascii=False, indent=2)
        log.info(f"  Quality report saved: {report_path}")

    return results


if __name__ == "__main__":
    main()
