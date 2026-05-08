"""Re-render cover letter PDF + DOCX with a new date, without re-running the LLM.

By default this reads eligible jobs from the Google Sheet, which is the source
of truth for the application pipeline. Use --source local only for offline
legacy runs against .tmp/scored_jobs.json.

Usage:
    python -m execution.update_cover_letter_dates --letter-date 2026-04-22 --min-score 6
    python -m execution.update_cover_letter_dates --letter-date 2026-04-22 --min-score 6 --source local
"""
from __future__ import annotations

import argparse
import json
import logging
import os
import re
import sys
from datetime import date, datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from execution.generate_cover_letter import generate_docx, generate_pdf  # noqa: E402
from execution.utils import enforce_revops_subtitle, generate_job_id  # noqa: E402

logging.basicConfig(level=logging.INFO, format="%(message)s")
log = logging.getLogger(__name__)

ROOT = Path(__file__).resolve().parent.parent
DEFAULT_SHEET_NAME = "Swiss Job Search Pipeline"
SCORED_JOBS = ROOT / ".tmp" / "scored_jobs.json"
CL_CHECKPOINT = ROOT / ".tmp" / "cover_letter_checkpoint.json"
APPS_DIR = ROOT / ".tmp" / "applications"
COVER_LETTER_DOCX = "Cover_Letter_Simon_Oberpertinger_Mair.docx"
COVER_LETTER_PDF = "Cover_Letter_Simon_Oberpertinger_Mair.pdf"
VALID_LANGUAGES = {"de", "en", "it"}

# Greeting patterns, inverse of _greeting_for_contact in generate_cover_letter.py
GREETING_PATTERNS = [
    (re.compile(r"^Guten Tag\s+(.+?),?\s*$", re.IGNORECASE), "de"),
    (re.compile(r"^Dear\s+(.+?),?\s*$", re.IGNORECASE), "en"),
    (re.compile(r"^Gentile\s+(.+?),?\s*$", re.IGNORECASE), "it"),
]
ANON_GREETINGS = {
    "sehr geehrte damen und herren,",
    "dear hiring manager,",
    "gentili signore e signori,",
}

SIGNOFF_MARKERS = {
    "freundliche grusse",
    "freundliche gruesse",
    "freundliche gr\u00fcsse",
    "mit freundlichen grussen",
    "mit freundlichen gruessen",
    "mit freundlichen gr\u00fcssen",
    "kind regards",
    "best regards",
    "cordiali saluti",
}


def _parse_contact_from_greeting(greeting: str) -> str | None:
    """Return the contact name from the greeting, or None if anonymous."""
    g = greeting.strip()
    if g.lower() in ANON_GREETINGS:
        return None
    if g.lower() in {"dear hiring manager,", "dear hiring manager"}:
        return None

    for pat, _lang in GREETING_PATTERNS:
        m = pat.match(g)
        if not m:
            continue
        name = m.group(1).strip().rstrip(",")
        if name.lower() in {"hiring manager", "sir or madam", "damen und herren"}:
            return None
        return name
    return None


def extract_body_and_contact(docx_path: Path) -> tuple[str, str | None]:
    """Return (body_text, contact_name) parsed from an existing cover-letter DOCX."""
    doc = Document(str(docx_path))
    paragraphs = [p.text for p in doc.paragraphs]

    greeting_idx = None
    for i, text in enumerate(paragraphs):
        t = text.strip()
        if not t:
            continue
        low = t.lower()
        if low in ANON_GREETINGS or any(low.startswith(p) for p in ("guten tag ", "dear ", "gentile ")):
            greeting_idx = i
            break

    if greeting_idx is None:
        # Fallback: skip header lines (name, subtitle, date, address — first ~10 non-empty
        # paragraphs) and treat the rest up to the signoff as the body. Handles edge-case
        # DOCX files where the greeting paragraph has unexpected formatting.
        log.warning(f"[WARN] Could not find greeting in {docx_path.parent.name} — using fallback extraction")
        non_empty = [(i, p.strip()) for i, p in enumerate(paragraphs) if p.strip()]
        # Header is typically: name, subtitle, date, blank lines, recipient block (~8-10 lines)
        body_start = min(10, len(non_empty) // 3) if len(non_empty) > 12 else 0
        greeting_idx = non_empty[body_start][0] if non_empty else 0
        contact_name = None
    else:
        contact_name = _parse_contact_from_greeting(paragraphs[greeting_idx])

    signoff_idx = None
    for i in range(greeting_idx + 1, len(paragraphs)):
        low = paragraphs[i].strip().lower().rstrip(",")
        if low in SIGNOFF_MARKERS:
            signoff_idx = i
            break
    if signoff_idx is None:
        raise ValueError(f"Could not find signoff in {docx_path}")

    body_paragraphs = [p.strip() for p in paragraphs[greeting_idx + 1:signoff_idx] if p.strip()]
    return "\n\n".join(body_paragraphs), contact_name


def _detect_language_from_greeting(greeting: str) -> str:
    """Return language code from the greeting line, defaulting to 'de'."""
    g = greeting.strip().lower()
    if g.startswith("dear ") or g.startswith("to whom"):
        return "en"
    if g.startswith("gentile ") or g.startswith("gentili "):
        return "it"
    return "de"


def _normalize_header(header: str) -> str:
    return header.strip().lower().replace(" ", "_")


def _safe_cell(row: list[str], col_map: dict[str, int], name: str) -> str:
    idx = col_map.get(name)
    if idx is None or idx >= len(row):
        return ""
    return row[idx].strip()


def _parse_score(raw: object) -> int:
    try:
        return int(float(str(raw).strip()))
    except (TypeError, ValueError):
        return 0


def _load_jobs_from_local(min_score: int) -> list[dict]:
    with open(SCORED_JOBS, "r", encoding="utf-8") as f:
        scored = json.load(f)

    jobs = []
    for job in scored:
        if _parse_score(job.get("score", 0)) < min_score:
            continue
        if not job.get("job_id"):
            job["job_id"] = generate_job_id(
                str(job.get("title", "")),
                str(job.get("company", "")),
                str(job.get("url", "")),
            )
        jobs.append(job)
    log.info(f"Local jobs with score >= {min_score}: {len(jobs)}")
    return jobs


def _load_jobs_from_sheet(sheet_name: str, min_score: int) -> list[dict]:
    """Read eligible jobs from the Google Sheet, the application source of truth."""
    from execution.write_jobs_to_sheet import authenticate, _sheets_api_call  # Lazy import keeps local mode offline.

    try:
        client = authenticate()
        worksheet = _sheets_api_call(client.open, sheet_name).sheet1
        rows = _sheets_api_call(worksheet.get_all_values)
    except Exception as e:
        raise RuntimeError(
            f"Could not read Google Sheet '{sheet_name}'. "
            "Use --source local for offline JSON mode."
        ) from e

    if not rows or len(rows) < 2:
        log.warning(f"Sheet '{sheet_name}' is empty or has no data rows")
        return []

    col_map = {_normalize_header(h): i for i, h in enumerate(rows[0])}
    required = ["score", "job_id", "title", "company", "location"]
    missing = [name for name in required if name not in col_map]
    if missing:
        available = sorted(col_map.keys())
        raise RuntimeError(
            f"Sheet '{sheet_name}' is missing required column(s): {', '.join(missing)}\n"
            f"Headers currently in the sheet (normalized): {', '.join(available)}\n"
            f"This usually means a column was renamed in the Sheet UI. Either rename "
            f"it back, or update the `required` list in _load_jobs_from_sheet() to "
            f"match the new header name."
        )

    has_cl_flag = "cl_generated" in col_map

    jobs = []
    skipped_no_cl = 0
    for row_idx, row in enumerate(rows[1:], start=2):
        score = _parse_score(_safe_cell(row, col_map, "score"))
        if score < min_score:
            continue

        # Only re-stamp rows whose CL has actually been generated. Score>=6 rows
        # added in the most recent pipeline_full but not yet in the CL stage
        # would otherwise be counted as "missing folder" and trip the coverage
        # alert in main() with a false positive.
        if has_cl_flag:
            cl_done = _safe_cell(row, col_map, "cl_generated").strip().lower()
            if cl_done != "yes":
                skipped_no_cl += 1
                continue

        job_id = _safe_cell(row, col_map, "job_id")
        if not job_id:
            log.warning(f"Skipping sheet row {row_idx}: missing Job_ID")
            continue

        jobs.append({
            "row_idx": row_idx,
            "score": score,
            "job_id": job_id,
            "title": _safe_cell(row, col_map, "title"),
            "company": _safe_cell(row, col_map, "company"),
            "location": _safe_cell(row, col_map, "location"),
            "url": _safe_cell(row, col_map, "url"),
            "description": _safe_cell(row, col_map, "description"),
        })

    log.info(
        f"Sheet rows with score >= {min_score} and CL_Generated=Yes: {len(jobs)}"
        + (f" (skipped {skipped_no_cl} not-yet-generated)" if has_cl_flag else "")
    )
    return jobs


def _load_jobs(source: str, sheet_name: str, min_score: int) -> list[dict]:
    if source == "sheet":
        return _load_jobs_from_sheet(sheet_name, min_score)
    if source == "local":
        return _load_jobs_from_local(min_score)
    raise ValueError(f"Unsupported source: {source}")


def _index_application_folders() -> dict[str, Path]:
    if not APPS_DIR.exists():
        return {}
    folders: dict[str, Path] = {}
    for folder in APPS_DIR.iterdir():
        if not folder.is_dir() or not folder.name.startswith("J-"):
            continue
        job_id = folder.name.split("_")[0]
        if job_id in folders:
            log.warning(f"Duplicate application folder for {job_id}: keeping {folders[job_id]}, ignoring {folder}")
            continue
        folders[job_id] = folder
    return folders


def _load_checkpoint() -> dict:
    if not CL_CHECKPOINT.exists():
        return {}
    try:
        with open(CL_CHECKPOINT, "r", encoding="utf-8") as f:
            return json.load(f).get("processed", {})
    except Exception as e:
        raise RuntimeError(f"Could not read cover letter checkpoint at {CL_CHECKPOINT}: {e}") from e


def _detect_language(docx_path: Path, checkpoint_entry: dict) -> str:
    checkpoint_language = checkpoint_entry.get("language")
    if checkpoint_language in VALID_LANGUAGES:
        return checkpoint_language

    doc = Document(str(docx_path))
    first_greeting = next(
        (
            p.text
            for p in doc.paragraphs
            if p.text.strip().lower().startswith(("guten tag", "dear ", "gentile ", "sehr geehrte", "to whom"))
        ),
        "",
    )
    return _detect_language_from_greeting(first_greeting)


def _document_mtime(fld: Path, docx_path: Path) -> date:
    pdf_path = fld / COVER_LETTER_PDF
    mtimes = [docx_path.stat().st_mtime]
    if pdf_path.exists():
        mtimes.append(pdf_path.stat().st_mtime)
    return datetime.fromtimestamp(max(mtimes)).date()


def _temp_path(path: Path) -> Path:
    return path.with_name(f"{path.stem}.updating{path.suffix}")


def _cleanup_temp_files(*paths: Path) -> None:
    for path in paths:
        try:
            path.unlink(missing_ok=True)
        except Exception as e:
            log.warning(f"Could not remove temporary file {path}: {e}")


def _replace_pair_transactionally(tmp_pdf: Path, pdf_path: Path, tmp_docx: Path, docx_path: Path) -> None:
    """Replace PDF and DOCX as a pair, rolling back if either replace fails."""
    backup_pdf = _temp_path(pdf_path.with_suffix(".rollback.pdf"))
    backup_docx = _temp_path(docx_path.with_suffix(".rollback.docx"))
    backups: list[tuple[Path, Path]] = []

    _cleanup_temp_files(backup_pdf, backup_docx)

    try:
        if pdf_path.exists():
            pdf_path.replace(backup_pdf)
            backups.append((backup_pdf, pdf_path))
        if docx_path.exists():
            docx_path.replace(backup_docx)
            backups.append((backup_docx, docx_path))

        tmp_pdf.replace(pdf_path)
        tmp_docx.replace(docx_path)

        _cleanup_temp_files(backup_pdf, backup_docx)
    except Exception:
        for backup, original in reversed(backups):
            if backup.exists():
                try:
                    backup.replace(original)
                except Exception as rollback_error:
                    log.error(f"Rollback failed for {original}: {rollback_error}")
        raise


# Kept for execution/update_phone.py which still calls this for full re-renders
# after a phone-number swap. The daily date refresh uses _restamp_atomically.
def _render_atomically(
    body: str,
    job: dict,
    lang_code: str,
    pdf_path: Path,
    docx_path: Path,
    contact_name: str | None,
    subtitle: str,
    letter_date: str,
) -> None:
    tmp_pdf = _temp_path(pdf_path)
    tmp_docx = _temp_path(docx_path)
    _cleanup_temp_files(tmp_pdf, tmp_docx)

    try:
        generate_pdf(
            text=body,
            job=job,
            language_code=lang_code,
            output_path=tmp_pdf,
            contact_name=contact_name,
            subtitle=subtitle,
            letter_date=letter_date,
        )
        generate_docx(
            text=body,
            job=job,
            language_code=lang_code,
            output_path=tmp_docx,
            contact_name=contact_name,
            subtitle=subtitle,
            letter_date=letter_date,
        )
        _replace_pair_transactionally(tmp_pdf, pdf_path, tmp_docx, docx_path)
    except Exception:
        _cleanup_temp_files(tmp_pdf, tmp_docx)
        raise


DATE_RX = re.compile(r"^\d{2}\.\d{2}\.\d{4}$")
LETTER_META = "letter_meta.json"


def _load_letter_meta(folder: Path) -> dict | None:
    """Read sidecar body+meta written at generation time. None if missing/corrupt."""
    meta = folder / LETTER_META
    if not meta.exists():
        return None
    try:
        return json.loads(meta.read_text(encoding="utf-8"))
    except Exception as e:
        log.warning(f"[meta] {folder.name}: corrupt {LETTER_META} - {e}")
        return None


def _save_letter_meta(
    folder: Path,
    body: str,
    contact_name: str | None,
    subtitle: str,
    language: str,
    first_letter_date: str | None,
) -> None:
    target = folder / LETTER_META
    payload = json.dumps(
        {
            "body": body,
            "contact_name": contact_name,
            "subtitle": subtitle,
            "language": language,
            "first_letter_date": first_letter_date,
        },
        ensure_ascii=False,
        indent=2,
    )
    tmp = target.with_suffix(target.suffix + ".part")
    tmp.write_text(payload, encoding="utf-8")
    tmp.replace(target)


def _mutate_docx_date_to_tmp(docx_path: Path, new_date_de: str, tmp_docx: Path) -> bool:
    """Open the existing DOCX, replace the right-aligned date paragraph in memory,
    save to `tmp_docx`. Returns True iff a matching paragraph was found and rewritten.

    Invariant: in the cover-letter template (generate_cover_letter.py:785) the date
    paragraph is the only right-aligned paragraph, so the first match is always it.
    """
    doc = Document(str(docx_path))
    for p in doc.paragraphs:
        if p.alignment != WD_ALIGN_PARAGRAPH.RIGHT:
            continue
        if not DATE_RX.match(p.text.strip()):
            continue
        # Generator always produces a single run with explicit Pt(10) sizing.
        # Mutating runs[0].text preserves <w:rPr> (font size, color); any extra
        # runs on a hand-edited DOCX are removed to avoid empty <w:r> cruft.
        if not p.runs:
            p.add_run(new_date_de)
        else:
            p.runs[0].text = new_date_de
            for run in list(p.runs[1:]):
                run._element.getparent().remove(run._element)
        doc.save(str(tmp_docx))
        return True
    return False


def _restamp_atomically(
    folder: Path,
    job: dict,
    pdf_path: Path,
    docx_path: Path,
    letter_date: str,
    fallback_checkpoint_entry: dict | None = None,
) -> None:
    """Refresh DOCX (in-place edit) + PDF (full re-render) for one folder.

    Reads body/contact/subtitle/language from letter_meta.json; falls back to regex
    extraction (and backfills the sidecar) for legacy folders. Both files are
    swapped together via _replace_pair_transactionally.
    """
    meta = _load_letter_meta(folder)
    if meta is None:
        body, contact_name = extract_body_and_contact(docx_path)
        if not body:
            raise RuntimeError("empty body after extraction")
        ck = fallback_checkpoint_entry or {}
        subtitle = enforce_revops_subtitle(ck.get("subtitle") or _read_existing_subtitle(docx_path))
        language = _detect_language(docx_path, ck)
        # Pass None for first_letter_date — we don't know the original generation
        # date for legacy folders, and the field is informational only.
        _save_letter_meta(folder, body, contact_name, subtitle, language, None)
        meta = {
            "body": body,
            "contact_name": contact_name,
            "subtitle": subtitle,
            "language": language,
        }
    else:
        # Re-enforce subtitle policy on cached value (cheap, idempotent).
        meta["subtitle"] = enforce_revops_subtitle(meta.get("subtitle"))

    new_date_de = datetime.strptime(letter_date, "%Y-%m-%d").strftime("%d.%m.%Y")

    tmp_pdf = _temp_path(pdf_path)
    tmp_docx = _temp_path(docx_path)
    _cleanup_temp_files(tmp_pdf, tmp_docx)

    try:
        if not _mutate_docx_date_to_tmp(docx_path, new_date_de, tmp_docx):
            raise RuntimeError("could not locate right-aligned date paragraph in DOCX")
        generate_pdf(
            text=meta["body"],
            job=job,
            language_code=meta["language"],
            output_path=tmp_pdf,
            contact_name=meta["contact_name"],
            subtitle=meta["subtitle"],
            letter_date=letter_date,
        )
        _replace_pair_transactionally(tmp_pdf, pdf_path, tmp_docx, docx_path)
    except Exception:
        _cleanup_temp_files(tmp_pdf, tmp_docx)
        raise


def _read_existing_subtitle(docx_path: Path) -> str | None:
    doc = Document(str(docx_path))
    if len(doc.paragraphs) < 2:
        return None
    subtitle = doc.paragraphs[1].text.strip()
    return subtitle or None


def _parse_since(value: str | None) -> date | None:
    if not value:
        return None
    try:
        return datetime.strptime(value, "%Y-%m-%d").date()
    except ValueError:
        log.error("Invalid --since, expected YYYY-MM-DD")
        sys.exit(1)


def _validate_letter_date(value: str) -> datetime:
    try:
        return datetime.strptime(value, "%Y-%m-%d")
    except ValueError:
        log.error("Invalid --letter-date, expected YYYY-MM-DD")
        sys.exit(1)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--letter-date", required=True, help="New letter date in YYYY-MM-DD format")
    ap.add_argument("--min-score", type=int, default=7)
    ap.add_argument("--sheet-name", default=DEFAULT_SHEET_NAME, help="Google Sheet name for --source sheet")
    ap.add_argument(
        "--source",
        choices=("sheet", "local"),
        default="sheet",
        help="Read eligible jobs from Google Sheet (default) or .tmp/scored_jobs.json",
    )
    ap.add_argument("--since", default=None, help="Only process folders modified on or after this date (YYYY-MM-DD)")
    ap.add_argument("--dry-run", action="store_true", help="Only parse/extract, do not rewrite files")
    ap.add_argument("--limit", type=int, default=None, help="Process only first N jobs (for testing)")
    ap.add_argument("--skip-drive", action="store_true", help="Skip Drive re-upload after rendering (fast mode)")
    args = ap.parse_args()

    parsed = _validate_letter_date(args.letter_date)
    log.info(f"Target letter date: {parsed.strftime('%d.%m.%Y')}")
    log.info(f"Job source: {args.source}")

    since_date = _parse_since(args.since)
    if since_date:
        log.info(f"Only processing folders modified on or after: {since_date}")

    try:
        high = _load_jobs(args.source, args.sheet_name, args.min_score)
    except Exception as e:
        log.error(e)
        sys.exit(1)

    try:
        checkpoint = _load_checkpoint()
    except Exception as e:
        log.error(e)
        sys.exit(1)
    all_folders = _index_application_folders()
    log.info(f"Application folders indexed: {len(all_folders)}")

    candidates = []
    missing_folder = 0
    missing_docx = 0
    for job in high:
        jid = job["job_id"]
        fld = all_folders.get(jid)
        if not fld:
            missing_folder += 1
            continue

        docx_path = fld / COVER_LETTER_DOCX
        if not docx_path.exists():
            missing_docx += 1
            continue

        if since_date:
            mtime = _document_mtime(fld, docx_path)
            if mtime < since_date:
                continue

        candidates.append(job)

    log.info(f"Candidates after folder/date filter: {len(candidates)}")
    log.info(f"Skipped missing folder: {missing_folder}; missing DOCX: {missing_docx}")

    # Capture candidate count before --limit truncation so coverage stays
    # meaningful when running with --limit for testing.
    matched_count = len(candidates)

    if args.limit:
        candidates = candidates[: args.limit]

    # Lazy import — keeps --source local fully offline. Imported once outside the
    # loop so a missing googleapiclient surfaces immediately, not per-row.
    upload_application_folder = None
    if not args.dry_run and not args.skip_drive:
        try:
            from execution.drive_upload import upload_application_folder as _uaf
            upload_application_folder = _uaf
        except Exception as e:
            log.warning(f"Drive upload disabled (import failed): {e}")
    elif args.skip_drive:
        log.info("Drive upload skipped (--skip-drive)")

    updated = 0
    errored = []
    for job in candidates:
        jid = job["job_id"]
        ck = checkpoint.get(jid, {})
        fld = all_folders[jid]
        docx_path = fld / COVER_LETTER_DOCX
        pdf_path = fld / COVER_LETTER_PDF

        if args.dry_run:
            meta = _load_letter_meta(fld)
            source = "sidecar"
            if meta is None:
                try:
                    body, contact_name = extract_body_and_contact(docx_path)
                    if not body:
                        raise RuntimeError("empty body")
                    meta = {
                        "body": body,
                        "contact_name": contact_name,
                        "language": _detect_language(docx_path, ck),
                        "subtitle": enforce_revops_subtitle(
                            ck.get("subtitle") or _read_existing_subtitle(docx_path)
                        ),
                    }
                    source = "regex (legacy)"
                except Exception as e:
                    log.error(f"[ERR ] {jid}: body extraction failed - {e}")
                    errored.append((jid, str(e)))
                    continue
            log.info(
                f"[DRY ] {jid} | score={job.get('score')} | lang={meta['language']} | "
                f"contact={meta['contact_name']!r} | body_words={len(meta['body'].split())} | "
                f"subtitle={meta['subtitle']!r} | source={source}"
            )
            updated += 1
            continue

        try:
            _restamp_atomically(
                folder=fld,
                job=job,
                pdf_path=pdf_path,
                docx_path=docx_path,
                letter_date=args.letter_date,
                fallback_checkpoint_entry=ck,
            )
            updated += 1
            log.info(f"[OK  ] {jid} | score={job.get('score')} | {fld.name[:80]}")

            if upload_application_folder is not None:
                try:
                    upload_application_folder(fld)
                except Exception as drive_err:
                    log.warning(f"[drive] {jid}: re-upload failed - {drive_err}")
        except Exception as e:
            log.error(f"[ERR ] {jid}: render failed - {e}")
            errored.append((jid, str(e)))

    log.info("")
    log.info(f"Done. updated={updated}, errored={len(errored)}")

    eligible = len(high)
    log.info(
        f"Coverage: {matched_count}/{eligible} folders matched "
        f"({matched_count / max(eligible, 1):.1%})"
    )
    missing_pct = (missing_folder + missing_docx) / max(eligible, 1)
    if missing_pct > 0.05:
        log.warning(
            f"[ALERT] Missing artifacts for {missing_pct:.1%} of eligible Sheet rows "
            f"(missing_folder={missing_folder}, missing_docx={missing_docx}, "
            f"eligible={eligible}). Likely Drive<->Volume drift."
        )
        _send_coverage_alert(
            missing_pct=missing_pct,
            missing_folder=missing_folder,
            missing_docx=missing_docx,
            eligible=eligible,
            updated=updated,
        )

    if errored:
        log.info("Errors:")
        for jid, msg in errored:
            log.info(f"  {jid}: {msg}")
        # Only treat as failure if error rate exceeds 5% — isolated bad DOCX files
        # should not mark the whole run as failed in Task Scheduler.
        error_rate = len(errored) / max(updated + len(errored), 1)
        if error_rate > 0.05:
            sys.exit(1)


def _send_coverage_alert(
    *,
    missing_pct: float,
    missing_folder: int,
    missing_docx: int,
    eligible: int,
    updated: int,
) -> None:
    """Best-effort email alert when coverage gap exceeds threshold. Never raises."""
    to = os.environ.get("ALERT_EMAIL", "simonobemair@gmail.com").strip()
    if not to:
        return
    try:
        from execution.gmail_send import send_email
        subject = (
            f"[DOE] Cover letter date refresh — missing artifacts for "
            f"{missing_pct:.0%} of score>=6 rows"
        )
        body = (
            f"missing_folder = {missing_folder}\n"
            f"missing_docx   = {missing_docx}\n"
            f"eligible       = {eligible}\n"
            f"updated        = {updated}\n\n"
            f"Coverage gap likely indicates Drive<->Volume drift. "
            f"Check hydrate_volume_from_drive logs for download failures, "
            f"or verify the Sheet rows have matching DOE Applications/J-* folders.\n"
        )
        send_email(subject, body, to)
    except Exception as e:
        log.warning(f"[alert-email] failed: {e}")


if __name__ == "__main__":
    main()
