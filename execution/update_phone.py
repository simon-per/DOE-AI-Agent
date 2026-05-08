"""Refresh the phone number across ABOUTME.md and existing application folders.

Updates ABOUTME.md (source of truth), then walks `.tmp/applications/J-*/` and:
- Re-renders cover letter PDF + DOCX (no LLM, body extracted from existing DOCX).
- Replaces the phone string inside CV DOCX and ATS CV DOCX (text replacement only).
- CV PDFs are NEVER touched — re-rendering them deterministically would need data
  not in the cv_checkpoint, and Word/COM rendering does not match the fpdf2 layout.
  They are reported as stale; refresh by re-running stage 5.

By default only folders whose underlying job in `Swiss Job Search Pipeline` was
scraped on/after 2026-05-03 are processed (older folders are pre-cutoff drift
and shouldn't be touched). Pass `--all` to bypass that filter.

Usage:
    python -m execution.update_phone --phone "+41 79 123 45 67"
    python -m execution.update_phone --phone "+41 79 123 45 67" --dry-run --limit 1
    python -m execution.update_phone --phone "+41 79 123 45 67" --force --limit 5
"""
from __future__ import annotations

import argparse
import logging
import re
import sys
from datetime import date, datetime
from pathlib import Path

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from execution import profile_loader  # noqa: E402
from execution.update_cover_letter_dates import (  # noqa: E402
    COVER_LETTER_DOCX,
    COVER_LETTER_PDF,
    DEFAULT_SHEET_NAME,
    _document_mtime,
    _index_application_folders,
    _normalize_header,
    _read_existing_subtitle,
    _render_atomically,
    extract_body_and_contact,
)
from execution.utils import enforce_revops_subtitle  # noqa: E402

logging.basicConfig(level=logging.INFO, format="%(message)s")
log = logging.getLogger(__name__)

ROOT = Path(__file__).resolve().parent.parent
ABOUTME_PATH = ROOT / "ABOUTME.md"

CV_DOCX = "CV_Simon_Oberpertinger_Mair.docx"
CV_PDF = "CV_Simon_Oberpertinger_Mair.pdf"
CV_ATS_DOCX = "CV_Simon_Oberpertinger_Mair_ATS.docx"
CV_ATS_PDF = "CV_Simon_Oberpertinger_Mair_ATS.pdf"

CV_PAIRS = [(CV_DOCX, CV_PDF), (CV_ATS_DOCX, CV_ATS_PDF)]

DEFAULT_SCRAPED_SINCE = "2026-05-03"

PHONE_LINE_RE = re.compile(r'^(\s*phone:\s*)"([^"]*)"\s*$', re.MULTILINE)


def _read_aboutme_phone() -> tuple[str, str]:
    text = ABOUTME_PATH.read_text(encoding="utf-8")
    m = PHONE_LINE_RE.search(text)
    if not m:
        raise RuntimeError(f'Could not locate `phone: "..."` line in {ABOUTME_PATH}')
    return text, m.group(2)


def _write_aboutme_phone(original_text: str, new_phone: str) -> None:
    new_text, count = PHONE_LINE_RE.subn(
        lambda m: f'{m.group(1)}"{new_phone}"', original_text, count=1
    )
    if count != 1:
        raise RuntimeError("phone line replacement matched 0 lines after a successful read")
    tmp = ABOUTME_PATH.with_name(ABOUTME_PATH.name + ".updating")
    tmp.write_text(new_text, encoding="utf-8")
    tmp.replace(ABOUTME_PATH)


def _replace_in_docx(path: Path, old: str, new: str) -> bool:
    """Replace `old` with `new` in every paragraph and table cell of a DOCX.

    Returns True if any change was written. Falls back to paragraph-level text
    rewrite when the phone string is split across multiple runs.
    """
    doc = Document(str(path))
    changed = False

    def _paragraph_replace(p) -> bool:
        local = False
        for run in p.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
                local = True
        if local:
            return True
        joined = "".join(r.text for r in p.runs)
        if old in joined:
            new_joined = joined.replace(old, new)
            if p.runs:
                p.runs[0].text = new_joined
                for run in p.runs[1:]:
                    run.text = ""
            else:
                p.text = new_joined
            log.debug(f"  fallback paragraph rewrite in {path.name}")
            return True
        return False

    for p in doc.paragraphs:
        if _paragraph_replace(p):
            changed = True

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if _paragraph_replace(p):
                        changed = True

    if changed:
        tmp = path.with_name(path.stem + ".updating" + path.suffix)
        doc.save(str(tmp))
        tmp.replace(path)
    return changed


def _parse_date_arg(value: str | None, flag: str) -> date | None:
    if not value:
        return None
    try:
        return datetime.strptime(value, "%Y-%m-%d").date()
    except ValueError:
        log.error(f"Invalid {flag}, expected YYYY-MM-DD")
        sys.exit(1)


def _validate_phone(value: str) -> str:
    v = value.strip()
    if not v or not any(c.isdigit() for c in v):
        log.error("Invalid --phone: must be non-empty and contain at least one digit")
        sys.exit(1)
    return v


def _detect_language_from_docx(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    first = next(
        (
            p.text
            for p in doc.paragraphs
            if p.text.strip().lower().startswith(("guten tag", "dear ", "gentile ", "sehr geehrte", "to whom"))
        ),
        "",
    )
    g = first.strip().lower()
    if g.startswith("dear ") or g.startswith("to whom"):
        return "en"
    if g.startswith("gentile ") or g.startswith("gentili "):
        return "it"
    return "de"


def _load_scraped_dates(sheet_name: str) -> dict[str, date]:
    """Return {job_id: date_scraped} from the pipeline sheet.

    Fail-closed: any error raises RuntimeError. Caller must surface it.
    """
    from execution.write_jobs_to_sheet import authenticate

    client = authenticate()
    ws = client.open(sheet_name).sheet1
    rows = ws.get_all_values()
    if not rows or len(rows) < 2:
        raise RuntimeError(f"Sheet '{sheet_name}' has no data rows")

    header = {_normalize_header(h): i for i, h in enumerate(rows[0])}
    ji = header.get("job_id")
    di = header.get("date_scraped")
    if ji is None or di is None:
        raise RuntimeError(
            f"Sheet '{sheet_name}' missing required columns (Job_ID, Date Scraped)"
        )

    out: dict[str, date] = {}
    for row in rows[1:]:
        if len(row) <= max(ji, di):
            continue
        jid = row[ji].strip()
        raw = row[di].strip()
        if not jid or not raw:
            continue
        # Accept both ISO timestamps (2026-04-20T19:53:21+00:00) and YYYY-MM-DD.
        try:
            d = datetime.fromisoformat(raw.replace("Z", "+00:00")).date()
        except ValueError:
            try:
                d = datetime.strptime(raw[:10], "%Y-%m-%d").date()
            except ValueError:
                continue
        out[jid] = d
    return out


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--phone", required=True, help='New phone number, e.g. "+41 79 123 45 67"')
    ap.add_argument("--since", default=None, help="Only process folders whose docs were modified on or after YYYY-MM-DD (mtime filter)")
    ap.add_argument(
        "--scraped-since",
        default=DEFAULT_SCRAPED_SINCE,
        help=f"Only process folders whose job was scraped on/after YYYY-MM-DD (default {DEFAULT_SCRAPED_SINCE}). Folders not in the sheet are also skipped.",
    )
    ap.add_argument("--all", action="store_true", help="Bypass the scrape-date filter (touch every folder, including pre-cutoff ones)")
    ap.add_argument("--sheet-name", default=DEFAULT_SHEET_NAME, help="Pipeline sheet name")
    ap.add_argument("--limit", type=int, default=None, help="Process only first N folders")
    ap.add_argument("--dry-run", action="store_true", help="Preview changes, write nothing")
    ap.add_argument("--skip-aboutme", action="store_true", help="Refresh artifacts only; leave ABOUTME.md untouched")
    ap.add_argument("--force", action="store_true", help="Run even when ABOUTME.md already has --phone (re-sync drifted artifacts)")
    args = ap.parse_args()

    new_phone = _validate_phone(args.phone)
    since_date = _parse_date_arg(args.since, "--since")
    scraped_since = _parse_date_arg(args.scraped_since, "--scraped-since") if not args.all else None

    aboutme_text, old_phone = _read_aboutme_phone()
    already_in_sync = old_phone == new_phone
    if already_in_sync and not args.force:
        log.info(f"Phone already set to {new_phone!r} in ABOUTME.md — nothing to do. Pass --force to re-sync artifacts.")
        return

    log.info(f"Old phone: {old_phone!r}")
    log.info(f"New phone: {new_phone!r}")
    if already_in_sync:
        log.info("ABOUTME.md already in sync; --force passed → re-sync artifacts only.")

    if not args.skip_aboutme and not already_in_sync:
        if args.dry_run:
            log.info("[DRY ] would rewrite ABOUTME.md phone line")
        else:
            _write_aboutme_phone(aboutme_text, new_phone)
            log.info("[OK  ] ABOUTME.md updated")

    profile_loader._cached_profile = None  # force fresh read on next render

    scraped_map: dict[str, date] = {}
    if scraped_since is not None:
        try:
            scraped_map = _load_scraped_dates(args.sheet_name)
            log.info(f"Sheet rows with Date Scraped: {len(scraped_map)}")
        except Exception as e:
            log.error(f"Could not load scrape-date map from '{args.sheet_name}': {e}")
            log.error("Aborting (fail-closed). Pass --all to bypass the scrape-date filter.")
            sys.exit(1)

    folders = _index_application_folders()
    log.info(f"Application folders indexed: {len(folders)}")

    candidates: list[Path] = []
    skipped_pre_cutoff = 0
    skipped_not_in_sheet = 0
    for fld in folders.values():
        jid = fld.name.split("_", 1)[0]

        if scraped_since is not None:
            scraped = scraped_map.get(jid)
            if scraped is None:
                skipped_not_in_sheet += 1
                continue
            if scraped < scraped_since:
                skipped_pre_cutoff += 1
                continue

        if since_date:
            cl_docx = fld / COVER_LETTER_DOCX
            if not cl_docx.exists():
                continue
            if _document_mtime(fld, cl_docx) < since_date:
                continue

        candidates.append(fld)

    candidates.sort(key=lambda p: p.name)
    if args.limit:
        candidates = candidates[: args.limit]

    if scraped_since is not None:
        log.info(
            f"Scrape filter ≥ {scraped_since}: eligible={len(candidates)}, "
            f"skipped_pre_cutoff={skipped_pre_cutoff}, skipped_not_in_sheet={skipped_not_in_sheet}"
        )
    log.info(f"Folders to process: {len(candidates)}")

    today_iso = datetime.now().strftime("%Y-%m-%d")
    cover_letters_updated = 0
    cv_docx_updated = 0
    stale_pdfs: list[Path] = []
    errored: list[tuple[str, str]] = []

    for fld in candidates:
        jid = fld.name.split("_", 1)[0]
        cl_docx = fld / COVER_LETTER_DOCX
        cl_pdf = fld / COVER_LETTER_PDF

        if cl_docx.exists():
            try:
                body, contact_name = extract_body_and_contact(cl_docx)
                subtitle = enforce_revops_subtitle(_read_existing_subtitle(cl_docx))
                lang_code = _detect_language_from_docx(cl_docx)
                job_stub = {"title": "", "company": "", "location": "", "url": ""}

                if args.dry_run:
                    log.info(f"[DRY ] {jid} cover letter | lang={lang_code} | contact={contact_name!r}")
                else:
                    _render_atomically(
                        body=body,
                        job=job_stub,
                        lang_code=lang_code,
                        pdf_path=cl_pdf,
                        docx_path=cl_docx,
                        contact_name=contact_name,
                        subtitle=subtitle,
                        letter_date=today_iso,
                    )
                    log.info(f"[OK  ] {jid} cover letter re-rendered")
                cover_letters_updated += 1
            except Exception as e:
                log.error(f"[ERR ] {jid} cover letter: {e}")
                errored.append((jid, f"cover letter: {e}"))

        for cv_name, pdf_name in CV_PAIRS:
            cv_path = fld / cv_name
            pdf_path = fld / pdf_name
            if not cv_path.exists():
                if pdf_path.exists():
                    stale_pdfs.append(pdf_path)
                continue
            try:
                if args.dry_run:
                    doc = Document(str(cv_path))
                    seen = any(old_phone in p.text for p in doc.paragraphs) or any(
                        old_phone in p.text
                        for table in doc.tables
                        for row in table.rows
                        for cell in row.cells
                        for p in cell.paragraphs
                    )
                    log.info(f"[DRY ] {jid} {cv_name} | phone present={seen}")
                    if seen:
                        cv_docx_updated += 1
                else:
                    if _replace_in_docx(cv_path, old_phone, new_phone):
                        cv_docx_updated += 1
                        log.info(f"[OK  ] {jid} {cv_name} phone replaced")
                    else:
                        log.warning(f"[skip] {jid} {cv_name}: old phone not found")
            except Exception as e:
                log.error(f"[ERR ] {jid} {cv_name}: {e}")
                errored.append((jid, f"{cv_name}: {e}"))

            if pdf_path.exists():
                stale_pdfs.append(pdf_path)

    log.info("")
    log.info(
        f"Done. cover_letters={cover_letters_updated}, cv_docx={cv_docx_updated}, "
        f"stale_cv_pdfs={len(stale_pdfs)}, errored={len(errored)}"
    )
    if stale_pdfs:
        log.info(
            f"Stale CV PDFs: {len(stale_pdfs)} files. "
            "Re-run stage 5 (`python -m execution.generate_cv`) to refresh them."
        )
    if errored:
        log.info("Errors:")
        for jid, msg in errored:
            log.info(f"  {jid}: {msg}")
        sys.exit(1)


if __name__ == "__main__":
    main()
