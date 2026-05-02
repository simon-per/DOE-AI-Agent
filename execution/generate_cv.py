"""
Stage 5: Generate tailored CVs (PDF + DOCX) for high-scoring jobs.
Uses HTML/CSS template rendered to PDF via Playwright (headless Chromium).
Swiss CV format: photo, personal info, work permit status.

Primary: OpenRouter (Qwen3 235B) - reliable, paid (low cost, no rate limits)
Fallback: Google AI Studio (Gemini 3 Flash) - free but rate-limited

Input:  .tmp/scored_jobs.json
Output: .tmp/applications/{company}_{title}/CV_Simon_Oberpertinger_Mair_{company}.pdf + .docx

Usage:
    python execution/generate_cv.py
    python execution/generate_cv.py --limit 2              # test with 2 jobs
    python execution/generate_cv.py --min-score 8           # only top matches
    python execution/generate_cv.py --reset-checkpoint      # re-generate all
"""

import argparse
import json
import logging
import os
import random
import re
import sys
import time
from pathlib import Path

import requests
from dotenv import load_dotenv
from jinja2 import Environment, FileSystemLoader
from playwright.sync_api import sync_playwright

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))
load_dotenv(PROJECT_ROOT / ".env")

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
log = logging.getLogger(__name__)

# Suppress noisy weasyprint/fontTools logs
logging.getLogger("weasyprint").setLevel(logging.WARNING)
logging.getLogger("fontTools").setLevel(logging.WARNING)

TMP_DIR = PROJECT_ROOT / ".tmp"
INPUT_FILE = TMP_DIR / "scored_jobs.json"
OUTPUT_DIR = TMP_DIR / "applications"
# Photo path: check env var first, fallback to default
_photo_env = os.getenv("CV_PHOTO_PATH", "").strip()
PHOTO_PATH = Path(_photo_env) if _photo_env else TMP_DIR / "CV_picture 24.06.2025_sizeAdjusted.jpg"
if not PHOTO_PATH.exists():
    log.warning(
        f"CV photo not found: {PHOTO_PATH}. "
        f"Set CV_PHOTO_PATH in .env or place file at expected path. "
        f"CVs will be generated without photo."
    )
TEMPLATE_DIR = Path(__file__).resolve().parent
CHECKPOINT_FILE = TMP_DIR / "cv_checkpoint.json"

DEFAULT_MIN_SCORE = 6

# Shared modules
from execution.llm_client import call_llm as _call_llm_shared, parse_json_response as _parse_json_response
from execution.language_detect import detect_language
from execution.utils import (
    sanitize_filename,
    clean_job_title,
    generate_job_id,
    enforce_revops_subtitle,
    find_application_folder,
    has_cv_outputs,
)
from execution.profile_loader import load_profile as _load_profile

CV_TAILOR_PROMPT = """You are tailoring a CV for a specific job application.

CANDIDATE PROFILE:
{profile}

JOB DETAILS:
Title: {title}
Company: {company}
Location: {location}
Key Matches: {key_matches}
Key Gaps: {key_gaps}

JOB DESCRIPTION (excerpt):
{description}

ANTI-HALLUCINATION RULES (CRITICAL):
- You may ONLY use numbers from the CANDIDATE PROFILE: 400,000+ CRM records, 1,000+ bulk uploads, 20% data integrity, 100+ CPQ attributes, 200+ material masters, 30-person team, 45 days notice, ~3 years experience.
- Do NOT invent percentages, revenue figures, time savings, or any other numbers.
- Do NOT use AI-giveaway words: results-driven, operational excellence, proven expertise, proven track record, leverage, utilize, streamline, robust, compelling, transformative, multifaceted, paramount, delve.

INSTRUCTIONS:
Write in {language_name}.

1. PROFESSIONAL SUMMARY: Write 3-4 sentences as a professional summary for this specific job.
   - Highlight the most relevant experience for THIS role
   - Mention key matching skills naturally
   - Modest but confident tone (Swiss style, NOT aggressive American self-promotion)
   - Do NOT mention lack of degree unless the job requires one
   - KEYWORD MIRRORING: Use exact phrases and terminology from the job description where they match the candidate's actual skills. For example, if the job says "Kundenstammdatenpflege" and the candidate manages CRM records, use that exact term. If the job says "Offertenwesen" and the candidate does CPQ, use that term. This helps ATS systems match the CV to the job.

2. SKILLS ORDER: List these skills in order of relevance to THIS job (most relevant first):
   CRM (Microsoft Dynamics), SAP MM/SD, Power BI, SQL, Excel, CPQ Configuration, Marketing Automation, Data Analysis, n8n & Make (Automation), API Integrations

3. SUBTITLE: Format must be exactly:
   "Revenue Operations Specialist | <3-5 skill keywords most relevant to this job>"
   - The prefix "Revenue Operations Specialist" is fixed — no translation, no synonyms (NOT "Spezialist", NOT the literal job title).
   - The suffix is 3-5 short skill keywords from THIS job's tech stack.
   - Use ONLY " • " (bullet with surrounding spaces) as the separator between keywords.
   - Each keyword is ONE concept. NEVER join two distinct tools with a space (e.g. "SAP CPQ" is WRONG because SAP and CPQ are two different products in this profile — write "SAP • CPQ" or "SAP S/4HANA • CPQ" instead). Same rule for any other pair like "Power BI Excel" or "Dynamics SAP".
   - Prefer generic category names (CRM, BI, CPQ, Automation, Analytics) for visual cleanness. EXCEPTION: keep specific product names when they are recognized industry-standard keywords that ATS systems explicitly filter for — specifically "Power BI", "Excel", "SAP", "Salesforce", "Dynamics". For this candidate, "Power BI" should always be written in full (not "BI") because it is a Microsoft-certified credential and a high-ROI brand keyword.
   Examples:
   - "Revenue Operations Specialist | CRM • Power BI • Automation"
   - "Revenue Operations Specialist | CRM • Power BI • CPQ • Automation"
   - "Revenue Operations Specialist | SAP S/4HANA • CPQ • Process Design"
   - "Revenue Operations Specialist | Dynamics 365 • Power BI • SQL"

CRITICAL: Respond with ONLY valid JSON:
{{"summary": "The 3-4 sentence summary here", "skills": ["Most relevant skill", "Second skill", ...], "subtitle": "Revenue Operations Specialist | <skill suffix>"}}"""


# ---------------------------------------------------------------------------
# LLM calling (same pattern as other scripts)
# ---------------------------------------------------------------------------

def _call_llm(openrouter_key: str | None, gemini_key: str | None, prompt: str) -> tuple[str, str]:
    """Call LLM via shared client. Uses max_tokens=800 and JSON mode for CV generation."""
    return _call_llm_shared(openrouter_key, gemini_key, prompt, temperature=0.3, max_tokens=800, json_mode=True)


def _get_profile(openrouter_key: str | None) -> str:
    profile = _load_profile()
    if openrouter_key:
        return profile.profile_anonymous()
    return profile.profile_full()

    # detect_language consolidated into execution/language_detect.py
    # _parse_json_response consolidated into execution/llm_client.py
    # sanitize_filename consolidated into execution/utils.py


# ---------------------------------------------------------------------------
# All profile data loaded from ABOUTME.md via profile_loader
# ---------------------------------------------------------------------------

def _get_labels(lang_code: str) -> dict[str, str]:
    return _load_profile().cv_labels(lang_code)

def _get_experience_bullets(lang_code: str) -> list[str]:
    return _load_profile().cv_experience_bullets(lang_code)

def _get_achievements(lang_code: str) -> list[dict]:
    return _load_profile().cv_achievements(lang_code)

def _get_interests(lang_code: str) -> list[str]:
    return _load_profile().cv_interests(lang_code)

def _get_language_list(lang_code: str) -> list[dict[str, str]]:
    return _load_profile().cv_language_list(lang_code)

def _get_bullet_keywords() -> dict[int, list[str]]:
    return _load_profile().experience.bullet_keywords

def _get_achievement_keywords() -> dict[int, list[str]]:
    return _load_profile().achievement_keywords


def _validate_skill_order(llm_skills: list[str], key_matches: list[str]) -> list[str]:
    """Ensure skills matching key_matches bubble to the top.

    [H6] Uses word-boundary matching (not substring) and deduplicates skills.
    Safety net: if LLM picks an odd order, deterministically re-sort.
    """
    if not key_matches:
        return llm_skills

    # Deduplicate skills (preserve first occurrence order)
    seen_lower = set()
    deduped = []
    for skill in llm_skills:
        if skill.lower() not in seen_lower:
            seen_lower.add(skill.lower())
            deduped.append(skill)
    if len(deduped) < len(llm_skills):
        log.info(f"  Removed {len(llm_skills) - len(deduped)} duplicate skills")
    llm_skills = deduped

    scored = []
    for skill in llm_skills:
        skill_lower = skill.lower()
        score = 0
        for km in key_matches:
            km_lower = km.lower()
            # Word-boundary matching in both directions (prevents "SAP" matching "ASAP")
            km_words = [w for w in km_lower.split() if len(w) > 2]
            skill_words = [w for w in skill_lower.split() if len(w) > 2]
            if any(re.search(rf'\b{re.escape(w)}\b', skill_lower) for w in km_words) or \
               any(re.search(rf'\b{re.escape(w)}\b', km_lower) for w in skill_words):
                score += 1
        scored.append((-score, skill))
    scored.sort()
    reordered = [s for _, s in scored]
    if reordered != llm_skills:
        log.info(f"  Skills reordered by key_matches: {reordered[:4]}")
    return reordered


COVER_LETTER_CHECKPOINT = Path(__file__).resolve().parent.parent / ".tmp" / "cover_letter_checkpoint.json"


def _get_cover_letter_subtitle(job: dict) -> str | None:
    """Read the cover letter subtitle from cover letter checkpoint for subtitle consistency."""
    entry = _get_cover_letter_checkpoint_entry(job)
    return entry.get("subtitle") or None if entry else None


def _get_cover_letter_language(job: dict) -> str | None:
    """Read detected language from cover letter checkpoint for CL/CV consistency."""
    entry = _get_cover_letter_checkpoint_entry(job)
    return entry.get("language") or None if entry else None


def _get_cover_letter_checkpoint_entry(job: dict) -> dict | None:
    """Read a job's entry from the cover letter checkpoint."""
    if not COVER_LETTER_CHECKPOINT.exists():
        return None
    try:
        with open(COVER_LETTER_CHECKPOINT, "r", encoding="utf-8") as f:
            cl_checkpoint = json.load(f)
        job_id = _get_job_id(job)
        return cl_checkpoint.get("processed", {}).get(job_id) or None
    except (json.JSONDecodeError, OSError):
        return None


def _reorder_by_relevance(items: list, keyword_map: dict, key_matches: list[str], description: str) -> list:
    """Reorder items by relevance to key_matches and job description.

    Scores each item by counting keyword hits in key_matches + description.
    Items with higher scores float to the top; original order is preserved
    for ties (stable sort).
    """
    match_text = " ".join(key_matches).lower() if key_matches else ""
    desc_lower = (description or "").lower()
    search_text = f"{match_text} {desc_lower}"

    scored = []
    for idx, item in enumerate(items):
        keywords = keyword_map.get(idx, [])
        score = sum(1 for kw in keywords if kw in search_text)
        scored.append((-score, idx, item))  # negative score for descending

    scored.sort()
    return [item for _, _, item in scored]


def _sort_projects_by_relevance(projects: list[dict], key_matches: list[str]) -> list[dict]:
    """Sort projects so the most tag-relevant project appears first.

    All projects are returned (unlike relevant_projects() which filters).
    Stable sort preserves ABOUTME.md order for ties.
    """
    km_lower = {w.lower() for w in key_matches}

    def score(proj: dict) -> int:
        return sum(1 for tag in proj.get("tags", []) if tag.lower() in km_lower)

    return sorted(projects, key=score, reverse=True)


# ---------------------------------------------------------------------------
# PDF generation
# ---------------------------------------------------------------------------

    # sanitize_filename consolidated into execution/utils.py (imported at top)


def generate_cv_pdf(
    summary: str,
    skills: list[str],
    subtitle: str,
    lang_code: str,
    output_path: Path,
    experience_bullets: list[str] | None = None,
    achievements: list[dict] | None = None,
    projects: list[dict] | None = None,
) -> Path:
    """Render the CV HTML template to PDF via Playwright (headless Chromium)."""
    env = Environment(loader=FileSystemLoader(str(TEMPLATE_DIR)))
    template = env.get_template("cv_template.html")

    # Resolve photo as file:// URI for the browser
    photo_uri = ""
    if PHOTO_PATH.exists():
        photo_uri = PHOTO_PATH.as_uri()
    else:
        log.warning(f"  CV photo not found: {PHOTO_PATH} — CV will be generated without photo")

    labels = _get_labels(lang_code)
    if experience_bullets is None:
        experience_bullets = _get_experience_bullets(lang_code)
    languages = _get_language_list(lang_code)
    if achievements is None:
        achievements = _get_achievements(lang_code)
    interests = _get_interests(lang_code)

    html_content = template.render(
        lang_code=lang_code,
        photo_path=photo_uri,
        labels=labels,
        skills=skills,
        languages=languages,
        subtitle=subtitle,
        summary=summary,
        experience_bullets=experience_bullets,
        achievements=achievements,
        interests=interests,
        cv_dob=_load_profile().personal.dob,
        edu_start=_load_profile().education.start_year,
        edu_end=_load_profile().education.end_year,
        personal=_load_profile().personal,
        experience=_load_profile().experience,
        projects=projects if projects is not None else _load_profile().cv_projects(lang_code),
        certifications=_get_certifications(lang_code),
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Write temp HTML file so Chromium can resolve local file:// image URIs
    tmp_html = output_path.parent / "cv_render.html"
    tmp_html.write_text(html_content, encoding="utf-8")

    with sync_playwright() as p:
        browser = p.chromium.launch()
        try:
            page = browser.new_page()
            page.goto(tmp_html.as_uri())
            page.pdf(
                path=str(output_path),
                format="A4",
                margin={"top": "0", "right": "0", "bottom": "0", "left": "0"},
                print_background=True,
            )
        finally:
            browser.close()

    tmp_html.unlink(missing_ok=True)
    return output_path


def generate_cv_docx(
    summary: str,
    skills: list[str],
    subtitle: str,
    lang_code: str,
    output_path: Path,
    experience_bullets: list[str] | None = None,
    achievements: list[dict] | None = None,
    projects: list[dict] | None = None,
) -> Path:
    """Generate a CV as editable DOCX using python-docx."""
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    labels = _get_labels(lang_code)
    if experience_bullets is None:
        experience_bullets = _get_experience_bullets(lang_code)
    languages = _get_language_list(lang_code)
    if achievements is None:
        achievements = _get_achievements(lang_code)
    interests = _get_interests(lang_code)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # Main 2-column table (sidebar | content) — no borders
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(5.0)
    table.columns[1].width = Cm(13.0)

    # Remove all table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = borders.makeelement(qn(f'w:{border_name}'), {
            qn('w:val'): 'none', qn('w:sz'): '0', qn('w:space'): '0', qn('w:color'): 'auto'
        })
        borders.append(border)
    tblPr.append(borders)

    sidebar = table.cell(0, 0)
    main = table.cell(0, 1)

    # --- SIDEBAR ---
    def add_sidebar_heading(cell, text):
        p = cell.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x4A, 0x90, 0xD9)
        p.space_before = Pt(10)
        p.space_after = Pt(4)

    def add_sidebar_text(cell, label, value, bold_label=True):
        p = cell.add_paragraph()
        if bold_label:
            run = p.add_run(f"{label}: ")
            run.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        run = p.add_run(value)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        p.space_after = Pt(1)

    def add_hyperlink(paragraph, text, url, font_size=Pt(8), color=RGBColor(0x4A, 0x90, 0xD9)):
        """Add a clickable hyperlink to a DOCX paragraph."""
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        hyperlink = paragraph._element.makeelement(qn("w:hyperlink"), {qn("r:id"): r_id})
        new_run = paragraph._element.makeelement(qn("w:r"), {})
        rPr = paragraph._element.makeelement(qn("w:rPr"), {})
        sz = paragraph._element.makeelement(qn("w:sz"), {qn("w:val"): str(int(font_size.pt * 2))})
        c = paragraph._element.makeelement(qn("w:color"), {qn("w:val"): f"{color[0]:02x}{color[1]:02x}{color[2]:02x}"})
        u = paragraph._element.makeelement(qn("w:u"), {qn("w:val"): "none"})
        rPr.append(sz)
        rPr.append(c)
        rPr.append(u)
        new_run.append(rPr)
        t = paragraph._element.makeelement(qn("w:t"), {})
        t.text = text
        new_run.append(t)
        hyperlink.append(new_run)
        paragraph._element.append(hyperlink)

    # Photo placeholder (text note since DOCX table image is complex)
    if PHOTO_PATH.exists():
        p = sidebar.paragraphs[0]  # Use existing first paragraph
        try:
            run = p.add_run()
            run.add_picture(str(PHOTO_PATH), width=Cm(3.5))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            run = p.add_run("[Photo]")
            run.font.size = Pt(8)
    else:
        sidebar.paragraphs[0].text = ""

    # Personal info
    add_sidebar_heading(sidebar, labels["personal_info"])
    add_sidebar_text(sidebar, labels["date_of_birth"], _load_profile().personal.dob)
    add_sidebar_text(sidebar, labels["nationality"], labels["nationality_value"])
    add_sidebar_text(sidebar, labels["permit"], labels["permit_value"])
    _p = _load_profile().personal
    add_sidebar_text(sidebar, labels["email_label"], _p.email)
    add_sidebar_text(sidebar, labels["phone_label"], _p.phone)
    add_sidebar_text(sidebar, labels["location_label"], labels["location_value"])
    add_sidebar_text(sidebar, labels["availability_label"], labels["availability_value"])
    add_sidebar_text(sidebar, labels["license_label"], labels["license_value"])
    # LinkedIn (clickable)
    if _p.linkedin_handle:
        p = sidebar.add_paragraph()
        run = p.add_run("LinkedIn: ")
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        add_hyperlink(p, _p.linkedin_handle, f"https://{_p.linkedin}", font_size=Pt(8))
        p.space_after = Pt(1)
    # GitHub (clickable)
    if _p.github_handle:
        p = sidebar.add_paragraph()
        run = p.add_run("GitHub: ")
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        add_hyperlink(p, _p.github_handle, f"https://{_p.github}", font_size=Pt(8))
        p.space_after = Pt(1)

    # Skills
    add_sidebar_heading(sidebar, labels["skills_heading"])
    for skill in skills:
        p = sidebar.add_paragraph()
        run = p.add_run(f"• {skill}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        p.space_after = Pt(1)

    # Languages
    add_sidebar_heading(sidebar, labels["languages_heading"])
    for lang in languages:
        p = sidebar.add_paragraph()
        run = p.add_run(f"{lang['name']}: {lang['level']}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        p.space_after = Pt(1)

    # Interests
    add_sidebar_heading(sidebar, labels["interests_heading"])
    for interest in interests:
        p = sidebar.add_paragraph()
        run = p.add_run(f"• {interest}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        p.space_after = Pt(1)

    # --- MAIN CONTENT ---
    def add_main_heading(cell, text):
        p = cell.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1A, 0x27, 0x44)
        p.space_before = Pt(12)
        p.space_after = Pt(4)

    # Name + subtitle
    p = main.paragraphs[0]
    run = p.add_run(_load_profile().personal.name)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1A, 0x27, 0x44)
    p.space_after = Pt(0)

    p = main.add_paragraph()
    run = p.add_run(subtitle)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x4A, 0x90, 0xD9)
    p.space_after = Pt(8)

    # Summary
    add_main_heading(main, labels["summary_heading"])
    p = main.add_paragraph()
    run = p.add_run(summary)
    run.font.size = Pt(9)
    p.space_after = Pt(4)

    # Achievements
    add_main_heading(main, labels["achievements_heading"])
    for ach in achievements:
        p = main.add_paragraph()
        run = p.add_run(f"{ach['number']}  ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x4A, 0x90, 0xD9)
        run = p.add_run(ach['label'])
        run.font.size = Pt(9)
        p.space_after = Pt(2)

    # Experience
    add_main_heading(main, labels["experience_heading"])
    p = main.add_paragraph()
    _exp = _load_profile().experience
    run = p.add_run(_exp.role)
    run.bold = True
    run.font.size = Pt(10)
    p.space_after = Pt(0)

    official_line = _exp.official_line(lang_code)
    if official_line:
        p = main.add_paragraph()
        run = p.add_run(official_line)
        run.italic = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        p.space_after = Pt(0)

    p = main.add_paragraph()
    run = p.add_run(f"{_exp.company}, {_exp.company_location} | 04/2023 – {labels['present']}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x64, 0x64, 0x64)
    p.space_after = Pt(4)

    for bullet in experience_bullets:
        p = main.add_paragraph()
        run = p.add_run(f"• {bullet}")
        run.font.size = Pt(9)
        p.space_after = Pt(1)

    # Technical Projects (above Education — highlights tech/AI skills)
    docx_projects = projects if projects is not None else _load_profile().cv_projects(lang_code)
    if docx_projects:
        add_main_heading(main, labels.get("projects_heading", "Technical Projects"))
        for proj in docx_projects:
            p = main.add_paragraph()
            run = p.add_run(proj["name"])
            run.bold = True
            run.font.size = Pt(9)
            run = p.add_run(f"  |  {proj['tech']}")
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x64, 0x64, 0x64)
            p.space_after = Pt(0)
            if proj.get("description"):
                p = main.add_paragraph()
                run = p.add_run(proj["description"])
                run.font.size = Pt(7.5)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                p.space_after = Pt(0)
            p = main.add_paragraph()
            proj_href = proj["url"] if proj["url"].startswith("http") else f"https://{proj['url']}"
            add_hyperlink(p, proj["url"], proj_href, font_size=Pt(8))
            p.space_after = Pt(2)

    # Education
    add_main_heading(main, labels["education_heading"])
    p = main.add_paragraph()
    run = p.add_run(labels["matura_title"])
    run.bold = True
    run.font.size = Pt(10)
    p.space_after = Pt(0)

    p = main.add_paragraph()
    _edu = _load_profile().education
    run = p.add_run(f"{labels['matura_school']} | {_edu.start_year} – {_edu.end_year}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x64, 0x64, 0x64)
    p.space_after = Pt(4)

    # Certifications (from ABOUTME.md)
    add_main_heading(main, labels["certifications_heading"])
    for cert in _get_certifications(lang_code):
        p = main.add_paragraph()
        run = p.add_run(f"• {cert['name']}")
        run.font.size = Pt(9)
        detail_parts = []
        if cert.get("issuer"):
            detail_parts.append(cert["issuer"])
        if cert.get("score"):
            detail_parts.append(cert["score"])
        if detail_parts:
            run = p.add_run(f" ({' — '.join(detail_parts)})")
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x64, 0x64, 0x64)
        p.space_after = Pt(1)

    # Save DOCX
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


# ---------------------------------------------------------------------------
# ATS-friendly CV generation (single-column, no photo, no graphics)
# ---------------------------------------------------------------------------

def _get_certifications(lang_code: str) -> list[dict]:
    """Return certifications list from ABOUTME.md for template rendering."""
    return [
        {"name": c.name, "issuer": c.issuer, "score": c.score}
        for c in _load_profile().certifications
    ]


def generate_cv_ats_pdf(
    summary: str,
    skills: list[str],
    subtitle: str,
    lang_code: str,
    output_path: Path,
    experience_bullets: list[str] | None = None,
    achievements: list[dict] | None = None,
    projects: list[dict] | None = None,
) -> Path:
    """Render ATS-friendly CV (single-column, no graphics) to PDF via Playwright."""
    env = Environment(loader=FileSystemLoader(str(TEMPLATE_DIR)))
    template = env.get_template("cv_template_ats.html")

    labels = _get_labels(lang_code)
    if experience_bullets is None:
        experience_bullets = _get_experience_bullets(lang_code)
    languages = _get_language_list(lang_code)
    if achievements is None:
        achievements = _get_achievements(lang_code)
    interests = _get_interests(lang_code)

    html_content = template.render(
        lang_code=lang_code,
        labels=labels,
        skills=skills,
        languages=languages,
        subtitle=subtitle,
        summary=summary,
        experience_bullets=experience_bullets,
        achievements=achievements,
        interests=interests,
        certifications=_get_certifications(lang_code),
        cv_dob=_load_profile().personal.dob,
        edu_start=_load_profile().education.start_year,
        edu_end=_load_profile().education.end_year,
        personal=_load_profile().personal,
        experience=_load_profile().experience,
        projects=projects if projects is not None else _load_profile().cv_projects(lang_code),
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)

    tmp_html = output_path.parent / "cv_ats_render.html"
    tmp_html.write_text(html_content, encoding="utf-8")

    with sync_playwright() as p:
        browser = p.chromium.launch()
        try:
            page = browser.new_page()
            page.goto(tmp_html.as_uri())
            page.pdf(
                path=str(output_path),
                format="A4",
                margin={"top": "15mm", "right": "15mm", "bottom": "15mm", "left": "15mm"},
                print_background=False,
            )
        finally:
            browser.close()

    tmp_html.unlink(missing_ok=True)
    return output_path


def generate_cv_ats_docx(
    summary: str,
    skills: list[str],
    subtitle: str,
    lang_code: str,
    output_path: Path,
    experience_bullets: list[str] | None = None,
    achievements: list[dict] | None = None,
    projects: list[dict] | None = None,
) -> Path:
    """Generate ATS-friendly CV as DOCX (single column, no table, no graphics)."""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    labels = _get_labels(lang_code)
    if experience_bullets is None:
        experience_bullets = _get_experience_bullets(lang_code)
    languages = _get_language_list(lang_code)
    if achievements is None:
        achievements = _get_achievements(lang_code)
    interests = _get_interests(lang_code)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    _p = _load_profile().personal
    _exp = _load_profile().experience

    def add_heading(text):
        """Add an uppercase section heading with bottom border."""
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11)
        p.space_before = Pt(12)
        p.space_after = Pt(4)
        # Add bottom border
        from docx.oxml.ns import qn
        pPr = p._element.get_or_add_pPr()
        pBdr = pPr.makeelement(qn('w:pBdr'), {})
        bottom = pBdr.makeelement(qn('w:bottom'), {
            qn('w:val'): 'single', qn('w:sz'): '4',
            qn('w:space'): '1', qn('w:color'): '333333',
        })
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_text(text, size=Pt(9), bold=False, space_after=Pt(2)):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = size
        run.bold = bold
        p.space_after = space_after
        return p

    # --- Name + Subtitle ---
    p = doc.add_paragraph()
    run = p.add_run(_p.name)
    run.bold = True
    run.font.size = Pt(16)
    p.space_after = Pt(0)

    add_text(subtitle, size=Pt(10), space_after=Pt(4))

    # --- Contact info (plain text, no hyperlinks) ---
    contact_lines = [
        f"{_p.email} | {_p.phone} | {labels['location_value']}",
        f"LinkedIn: {_p.linkedin} | GitHub: {_p.github}",
        f"{labels['date_of_birth']}: {_p.dob} | {labels['nationality']}: {labels['nationality_value']} | {labels['permit']}: {labels['permit_value']}",
        f"{labels['availability_label']}: {labels['availability_value']} | {labels['license_label']}: {labels['license_value']}",
    ]
    for line in contact_lines:
        add_text(line, size=Pt(8.5), space_after=Pt(1))

    # --- Summary ---
    add_heading(labels["summary_heading"])
    add_text(summary, size=Pt(9.5), space_after=Pt(4))

    # --- Skills (comma-separated) ---
    add_heading(labels["skills_heading"])
    add_text(", ".join(skills), size=Pt(9), space_after=Pt(4))

    # --- Achievements ---
    add_heading(labels["achievements_heading"])
    for ach in achievements:
        p = doc.add_paragraph()
        run = p.add_run(f"{ach['number']} ")
        run.bold = True
        run.font.size = Pt(9)
        run = p.add_run(f"— {ach['label']}")
        run.font.size = Pt(9)
        p.space_after = Pt(1)

    # --- Experience ---
    add_heading(labels["experience_heading"])
    p = doc.add_paragraph()
    run = p.add_run(_exp.role)
    run.bold = True
    run.font.size = Pt(10)
    official_line = _exp.official_line(lang_code)
    if official_line:
        # Inline on same paragraph for ATS parser compatibility (matches cv_template_ats.html)
        run = p.add_run(f" {official_line}")
        run.italic = True
        run.font.size = Pt(9.5)
    p.space_after = Pt(0)

    add_text(
        f"{_exp.company} — {_exp.company_location} | {_exp.role_start} – {labels['present']}",
        size=Pt(9), space_after=Pt(4),
    )

    for bullet in experience_bullets:
        p = doc.add_paragraph()
        run = p.add_run(f"• {bullet}")
        run.font.size = Pt(9)
        p.space_after = Pt(1)

    # --- Projects ---
    docx_projects = projects if projects is not None else _load_profile().cv_projects(lang_code)
    if docx_projects:
        add_heading(labels.get("projects_heading", "Technical Projects"))
        for proj in docx_projects:
            p = doc.add_paragraph()
            run = p.add_run(proj["name"])
            run.bold = True
            run.font.size = Pt(9)
            run = p.add_run(f" | {proj['tech']}")
            run.font.size = Pt(8)
            p.space_after = Pt(0)
            if proj.get("description"):
                add_text(proj["description"], size=Pt(8), space_after=Pt(0))
            add_text(proj["url"], size=Pt(8), space_after=Pt(2))

    # --- Education ---
    add_heading(labels["education_heading"])
    p = doc.add_paragraph()
    run = p.add_run(labels["matura_title"])
    run.bold = True
    run.font.size = Pt(10)
    p.space_after = Pt(0)
    add_text(
        f"{labels['matura_school']} | {_load_profile().education.start_year} – {_load_profile().education.end_year}",
        size=Pt(9), space_after=Pt(4),
    )

    # --- Certifications ---
    add_heading(labels["certifications_heading"])
    for cert in _get_certifications(lang_code):
        score_str = f" ({cert['score']})" if cert.get("score") else ""
        issuer_str = f" — {cert['issuer']}" if cert.get("issuer") else ""
        add_text(f"• {cert['name']}{issuer_str}{score_str}", size=Pt(9), space_after=Pt(1))

    # --- Languages ---
    add_heading(labels["languages_heading"])
    lang_str = ", ".join(f"{l['name']} ({l['level']})" for l in languages)
    add_text(lang_str, size=Pt(9), space_after=Pt(4))

    # --- Interests ---
    add_heading(labels["interests_heading"])
    add_text(", ".join(interests), size=Pt(9), space_after=Pt(4))

    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


# ---------------------------------------------------------------------------
# Checkpointing (resume on interruption)
# ---------------------------------------------------------------------------

def _get_job_id(job: dict) -> str:
    """Create a unique ID for a job based on title + company + URL."""
    return generate_job_id(job.get("title", ""), job.get("company", ""), job.get("url", ""))


def _load_checkpoint() -> dict:
    """Load checkpoint tracking which jobs have been processed."""
    if CHECKPOINT_FILE.exists():
        try:
            with open(CHECKPOINT_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except (json.JSONDecodeError, OSError):
            return {"processed": {}}
    return {"processed": {}}


def _save_checkpoint(checkpoint: dict):
    """Persist checkpoint to disk."""
    CHECKPOINT_FILE.parent.mkdir(parents=True, exist_ok=True)
    with open(CHECKPOINT_FILE, "w", encoding="utf-8") as f:
        json.dump(checkpoint, f, ensure_ascii=False, indent=2)


# ---------------------------------------------------------------------------
# Main pipeline
# ---------------------------------------------------------------------------

def generate_cv_for_job(
    openrouter_key: str | None,
    gemini_key: str | None,
    job: dict,
    skip_ats: bool = False,
    ats_only: bool = False,
) -> dict:
    """Generate a tailored CV for a single job."""
    title = job.get("title", "Unknown")
    company = job.get("company", "Unknown")
    location = job.get("location", "Unknown")
    description = job.get("description", "")
    key_matches = job.get("key_matches", [])
    key_gaps = job.get("key_gaps", [])

    # Detect recruitment agency — use actual employer for CV consistency with cover letter
    try:
        from execution.generate_cover_letter import _detect_actual_employer
        actual_employer = _detect_actual_employer(company, description)
        _strip = lambda s: re.sub(r'\.(com|ch|de|net|org)$', '', s.lower().strip()).strip()
        if actual_employer and _strip(actual_employer) != _strip(company):
            log.info(f"  Recruitment agency detected: '{company}' → actual employer: '{actual_employer}'")
            company = actual_employer
    except ImportError:
        pass

    # Prefer language from cover letter checkpoint (guarantees CL/CV consistency)
    cl_lang = _get_cover_letter_language(job)
    if cl_lang:
        lang_code = cl_lang
        lang_name = {"de": "German", "en": "English", "it": "Italian"}.get(cl_lang, cl_lang)
        log.info(f"  Language (from CL checkpoint): {lang_name}")
    else:
        lang_code, lang_name = detect_language(title, description)
        log.info(f"  Language (detected): {lang_name}")

    if len(description) > 2000:
        description = description[:2000] + "..."

    profile = _get_profile(openrouter_key)
    prompt = CV_TAILOR_PROMPT.format(
        profile=profile,
        title=title,
        company=company,
        location=location,
        key_matches=", ".join(key_matches) if key_matches else "General fit",
        key_gaps=", ".join(key_gaps) if key_gaps else "None identified",
        description=description,
        language_name=lang_name,
    )

    # Call LLM for tailored content
    for json_attempt in range(2):
        try:
            text, provider = _call_llm(openrouter_key, gemini_key, prompt)
            log.info(f"  Provider: {provider}")
            result = _parse_json_response(text)
            break
        except json.JSONDecodeError:
            if json_attempt == 0:
                log.warning(f"  JSON parse failed, retrying...")
                time.sleep(1)
                continue
            log.warning(f"  JSON parse failed twice, using defaults")
            result = {
                "summary": f"Experienced Revenue Operations Specialist with 3 years of CRM, SAP, CPQ, and revenue analytics expertise.",
                "skills": ["CRM (Microsoft Dynamics)", "SAP MM/SD", "Power BI", "SQL",
                           "Excel", "CPQ Configuration", "Marketing Automation", "Data Analysis"],
                "subtitle": enforce_revops_subtitle(None),
            }

    summary = result.get("summary", "")
    if not summary.strip():
        summary = "Experienced Revenue Operations Specialist with 3 years of CRM, SAP, CPQ, and revenue analytics expertise."
        log.warning("  Empty summary from LLM — using default")
    skills = result.get("skills", [])
    raw_subtitle = result.get("subtitle")
    subtitle = enforce_revops_subtitle(raw_subtitle)
    if raw_subtitle and raw_subtitle != subtitle:
        log.info(f"  Subtitle normalized: {raw_subtitle!r} → {subtitle!r}")

    # Safety net: ensure skills are ordered by relevance to key_matches
    skills = _validate_skill_order(skills, key_matches)

    # Subtitle consistency: prefer cover letter subtitle if available (re-normalized to catch stale checkpoints)
    cl_subtitle = _get_cover_letter_subtitle(job)
    if cl_subtitle:
        normalized_cl = enforce_revops_subtitle(cl_subtitle)
        if normalized_cl != subtitle:
            log.info(f"  Subtitle aligned with cover letter: '{normalized_cl}' (was: '{subtitle}')")
            subtitle = normalized_cl

    # Reorder experience bullets and achievements based on job relevance
    reordered_bullets = _reorder_by_relevance(
        _get_experience_bullets(lang_code),
        _get_bullet_keywords(), key_matches, description,
    )
    reordered_achievements = _reorder_by_relevance(
        _get_achievements(lang_code),
        _get_achievement_keywords(), key_matches, description,
    )
    log.info(f"  Bullets reordered: {[b[:30] + '...' for b in reordered_bullets[:3]]}")

    # Sort projects: most relevant to this job first (automation/CRM vs analytics)
    sorted_projects = _sort_projects_by_relevance(_load_profile().cv_projects(lang_code), key_matches)
    log.info(f"  Projects sorted: {[p['name'] for p in sorted_projects]}")

    # Generate outputs
    job_id = _get_job_id(job)
    safe_company = sanitize_filename(company) if company else "Unknown"
    safe_title = sanitize_filename(clean_job_title(title)) if title else "Unknown"
    folder_name = f"{job_id}_{safe_company}_{safe_title}"
    safe_name = sanitize_filename(_load_profile().personal.name)

    result = {
        "title": title,
        "company": company,
        "language": lang_name,
        "subtitle": subtitle,
    }

    render_args = dict(
        experience_bullets=reordered_bullets,
        achievements=reordered_achievements,
        projects=sorted_projects,
    )

    # Fancy CV (with photo, sidebar, colors)
    if not ats_only:
        output_path = OUTPUT_DIR / folder_name / f"CV_{safe_name}.pdf"
        pdf_path = generate_cv_pdf(
            summary, skills, subtitle, lang_code, output_path, **render_args,
        )
        log.info(f"  PDF saved: {pdf_path}")

        docx_output = output_path.with_suffix(".docx")
        docx_path = generate_cv_docx(
            summary, skills, subtitle, lang_code, docx_output, **render_args,
        )
        log.info(f"  DOCX saved: {docx_path}")
        result["pdf_path"] = str(pdf_path)
        result["docx_path"] = str(docx_path)

    # ATS-friendly CV (single-column, no photo, no graphics)
    if not skip_ats:
        ats_pdf_path = OUTPUT_DIR / folder_name / f"CV_{safe_name}_ATS.pdf"
        ats_pdf = generate_cv_ats_pdf(
            summary, skills, subtitle, lang_code, ats_pdf_path, **render_args,
        )
        log.info(f"  ATS PDF saved: {ats_pdf}")

        ats_docx_path = ats_pdf_path.with_suffix(".docx")
        ats_docx = generate_cv_ats_docx(
            summary, skills, subtitle, lang_code, ats_docx_path, **render_args,
        )
        log.info(f"  ATS DOCX saved: {ats_docx}")
        result["ats_pdf_path"] = str(ats_pdf)
        result["ats_docx_path"] = str(ats_docx)

    return result


def _load_applying_from_sheet(sheet_name: str = "Swiss Job Search Pipeline") -> list[dict]:
    """Read rows with Status=APPLYING and no CV yet from the Google Sheet.

    Returns job dicts compatible with generate_cv_for_job().
    Tries to enrich description from scored_jobs.json (full text first).
    """
    from execution.write_jobs_to_sheet import authenticate, _sheets_api_call

    client = authenticate()
    try:
        ws = _sheets_api_call(client.open, sheet_name).sheet1
    except Exception as e:
        log.error(f"[sheet-triggered] Cannot open sheet '{sheet_name}': {e}")
        return []

    rows = _sheets_api_call(ws.get_all_values)
    if not rows or len(rows) < 2:
        return []

    header = [h.strip() for h in rows[0]]

    def col(name: str) -> int | None:
        try:
            return header.index(name)
        except ValueError:
            return None

    ci = {name: col(name) for name in (
        "Job_ID", "Title", "Company", "Location", "Source",
        "Score", "Key Matches", "Key Gaps", "Description", "URL",
        "Date Posted", "Status", "CV_Generated", "Reasoning",
    )}

    # Load scored_jobs.json for full descriptions (best-effort)
    scored_by_id: dict[str, dict] = {}
    if INPUT_FILE.exists():
        try:
            with open(INPUT_FILE, "r", encoding="utf-8") as f:
                for j in json.load(f):
                    scored_by_id[_get_job_id(j)] = j
        except Exception:
            pass

    def cell(row: list, name: str) -> str:
        idx = ci.get(name)
        return row[idx].strip() if idx is not None and idx < len(row) else ""

    jobs = []
    for row in rows[1:]:
        status = cell(row, "Status").lower()
        cv_done = cell(row, "CV_Generated").strip().lower()
        if status != "applying" or cv_done == "yes":
            continue

        job_id = cell(row, "Job_ID")
        url    = cell(row, "URL")

        if job_id and job_id in scored_by_id:
            base = scored_by_id[job_id]
            description = base.get("description", cell(row, "Description"))
        else:
            description = cell(row, "Description")

        try:
            score = int(cell(row, "Score"))
        except ValueError:
            score = 0

        jobs.append({
            "job_id": job_id,
            "title": cell(row, "Title"),
            "company": cell(row, "Company"),
            "location": cell(row, "Location"),
            "source": cell(row, "Source"),
            "url": url,
            "description": description,
            "score": score,
            "key_matches": [k.strip() for k in cell(row, "Key Matches").split(",") if k.strip()],
            "key_gaps": [k.strip() for k in cell(row, "Key Gaps").split(",") if k.strip()],
            "reasoning": cell(row, "Reasoning"),
            "date_posted": cell(row, "Date Posted"),
        })

    log.info(f"[sheet-triggered] Found {len(jobs)} APPLYING rows without CV in '{sheet_name}'")
    return jobs


def main():
    parser = argparse.ArgumentParser(description="Generate tailored CVs for scored jobs")
    parser.add_argument("--limit", type=int, default=0, help="Limit number of CVs (0 = all)")
    parser.add_argument("--min-score", type=int, default=DEFAULT_MIN_SCORE, help="Minimum score threshold")
    parser.add_argument("--reset-checkpoint", action="store_true", help="Clear checkpoint and re-generate all")
    parser.add_argument("--cl-only", action="store_true", help="Only generate CVs for jobs with existing cover letters")
    parser.add_argument("--no-ats", action="store_true", help="Skip ATS-friendly CV generation")
    parser.add_argument("--ats-only", action="store_true", help="Only generate ATS-friendly CVs (skip fancy version)")
    parser.add_argument("--job-id", type=str, default=None, help="Generate only for a specific job_id (e.g. J-ab2e15)")
    parser.add_argument("--sheet-triggered", action="store_true",
                        help="Read APPLYING rows (with no CV yet) from Sheet instead of scored_jobs.json.")
    args = parser.parse_args()

    openrouter_key = os.getenv("OPEN_ROUTER_API_KEY")
    gemini_key = os.getenv("GOOGLE_AI_STUDIO_API_KEY")

    if not openrouter_key and not gemini_key:
        log.error("No API keys found. Set OPEN_ROUTER_API_KEY or GOOGLE_AI_STUDIO_API_KEY in .env")
        sys.exit(1)

    # Load or reset checkpoint
    if args.reset_checkpoint and CHECKPOINT_FILE.exists():
        CHECKPOINT_FILE.unlink()
        log.info("Checkpoint cleared — will re-generate all CVs")

    checkpoint = _load_checkpoint()

    # --sheet-triggered: read APPLYING rows with no CV from Sheet
    if args.sheet_triggered:
        scored_jobs = _load_applying_from_sheet()
        if not scored_jobs:
            log.info("[sheet-triggered] No APPLYING rows without CV. Nothing to do.")
            return
        log.info(f"[sheet-triggered] Processing {len(scored_jobs)} APPLYING job(s)")
        eligible_jobs = scored_jobs
    else:
        if not INPUT_FILE.exists():
            log.error(f"Input file not found: {INPUT_FILE}")
            log.error("Run execution/evaluate_jobs.py first (Stage 2)")
            sys.exit(1)

        with open(INPUT_FILE, "r", encoding="utf-8") as f:
            scored_jobs = json.load(f)

        log.info(f"Loaded {len(scored_jobs)} scored jobs")

        eligible_jobs = [j for j in scored_jobs if j.get("score", 0) >= args.min_score]
        log.info(f"Jobs scoring >= {args.min_score}: {len(eligible_jobs)}")

    if args.job_id:
        eligible_jobs = [j for j in eligible_jobs if _get_job_id(j) == args.job_id]
        if not eligible_jobs:
            log.error(f"Job ID {args.job_id} not found in scored_jobs.json (or below --min-score {args.min_score})")
            sys.exit(1)
        log.info(f"--job-id filter: targeting {args.job_id}")

    if args.cl_only:
        cl_checkpoint_file = TMP_DIR / "cover_letter_checkpoint.json"
        if cl_checkpoint_file.exists():
            with open(cl_checkpoint_file, "r", encoding="utf-8") as f:
                cl_job_ids = set(json.load(f).get("processed", {}).keys())
        else:
            cl_job_ids = set()
        before = len(eligible_jobs)
        eligible_jobs = [j for j in eligible_jobs if _get_job_id(j) in cl_job_ids]
        log.info(f"--cl-only: filtered to {len(eligible_jobs)} jobs with cover letters (from {before})")

    if args.limit > 0:
        eligible_jobs = eligible_jobs[:args.limit]
        log.info(f"Limiting to {len(eligible_jobs)} CVs")

    if not eligible_jobs:
        log.info("No jobs meet the score threshold.")
        return

    # Filter out jobs with empty company (can't produce professional CVs)
    no_company = [j for j in eligible_jobs if not j.get("company", "").strip()]
    if no_company:
        log.warning(f"Skipping {len(no_company)} jobs with no company name")
        for j in no_company:
            log.warning(f"  - '{j.get('title', '?')}' (score: {j.get('score', '?')})")
        eligible_jobs = [j for j in eligible_jobs if j.get("company", "").strip()]

    # Filter out already-processed jobs (checkpoint + folder-existence safety net)
    already_done = checkpoint.get("processed", {})
    pending_jobs = []
    folder_synced = 0
    for job in eligible_jobs:
        job_id = _get_job_id(job)
        if job_id in already_done:
            log.debug(f"  Skipping (checkpoint): {job.get('title', '?')} at {job.get('company', '?')}")
            continue
        existing_folder = find_application_folder(OUTPUT_DIR, job_id)
        if existing_folder and has_cv_outputs(existing_folder):
            log.info(f"  Skipping (folder exists): {job.get('title', '?')} at {existing_folder.name}")
            existing_pdf = next(
                (p for p in existing_folder.glob("CV_*.pdf") if "_ATS" not in p.name),
                None,
            )
            already_done[job_id] = {
                "title": job.get("title", ""),
                "company": job.get("company", ""),
                "pdf_path": str(existing_pdf) if existing_pdf else "",
            }
            folder_synced += 1
            continue
        pending_jobs.append(job)

    if folder_synced > 0:
        checkpoint["processed"] = already_done
        _save_checkpoint(checkpoint)
        log.info(f"Synced {folder_synced} pre-existing folder(s) to checkpoint (no LLM cost)")

    skipped = len(eligible_jobs) - len(pending_jobs)
    if skipped > 0:
        log.info(f"Skipping {skipped} already-generated CVs (checkpoint + folder dedup)")

    if not pending_jobs:
        log.info("All eligible jobs already have CVs. Use --reset-checkpoint to re-generate.")
        return

    log.info(f"Generating CVs for {len(pending_jobs)} jobs")

    results = []
    for i, job in enumerate(pending_jobs):
        log.info(f"[{i + 1}/{len(pending_jobs)}] Generating CV: {job.get('title', '?')} at {job.get('company', '?')} (score: {job.get('score', '?')})")
        try:
            result = generate_cv_for_job(
                openrouter_key, gemini_key, job,
                skip_ats=args.no_ats, ats_only=args.ats_only,
            )
            results.append(result)

            # Save to checkpoint
            job_id = _get_job_id(job)
            checkpoint.setdefault("processed", {})[job_id] = {
                "title": job.get("title", "?"),
                "company": job.get("company", "?"),
                "pdf_path": result.get("pdf_path", ""),
            }
            _save_checkpoint(checkpoint)

            # Update Google Sheet with CV status
            try:
                from execution.write_jobs_to_sheet import update_job_columns
                update_job_columns(
                    sheet_name="Swiss Job Search Pipeline",
                    job_url=job.get("url", ""),
                    updates={"CV_Generated": "Yes"},
                )
            except Exception as e:
                log.debug(f"  Sheet update skipped: {e}")

            # Upload application folder to Google Drive
            try:
                from execution.drive_upload import upload_application_folder
                pdf_path = result.get("pdf_path", "")
                if pdf_path:
                    drive_url = upload_application_folder(Path(pdf_path).parent)
                    if drive_url:
                        log.info(f"  Drive: {drive_url}")
            except Exception as e:
                log.debug(f"  Drive upload skipped: {e}")
        except Exception as e:
            log.error(f"  Failed: {e}")
            results.append({
                "title": job.get("title", "?"),
                "company": job.get("company", "?"),
                "error": str(e),
            })

        if i < len(pending_jobs) - 1:
            time.sleep(2.0)

    successful = [r for r in results if "pdf_path" in r]
    failed = [r for r in results if "error" in r]

    log.info(f"\n{'='*60}")
    log.info(f"CV Generation Complete")
    log.info(f"  Generated: {len(successful)}/{len(results)}")
    if skipped > 0:
        log.info(f"  Skipped (checkpoint): {skipped}")
    if failed:
        log.info(f"  Failed: {len(failed)}")
    log.info(f"  Output directory: {OUTPUT_DIR}")
    for r in successful:
        log.info(f"  - {r['company']}: {r['title']} ({r['language']}) [{r['subtitle']}]")

    return results


if __name__ == "__main__":
    main()
